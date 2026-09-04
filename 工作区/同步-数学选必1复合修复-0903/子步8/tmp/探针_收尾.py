# -*- coding: utf-8 -*-
"""T2实物探针·收尾轮：①H件题号「2.8.2-1」段 ②H件题号「2.8-1」段 ③X1首个题号段
＋全6讲练件 C9C9C9题号run全量扫描（判行4范围断言真伪）
＋I1/I2条目号与〔基〕〔进〕计数（核册目录页件级行条数）"""
import sys, io, json, re, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
from docx import Document
from docx.oxml.ns import qn

BASE = r'C:\提示词\高中数学\高中数学同步'
F = {
 'X1': '人教B版选必1 第1章 空间向量与立体几何·衔接件（29题）.docx',
 'I1': '人教B版选必1 第1章 空间向量与立体几何·知识清单（完成）.docx',
 'B': '人教B版选必1 第1章 空间向量与立体几何（上）·讲练件（61题）.docx',
 'C': '人教B版选必1 第1章 空间向量与立体几何（下）·讲练件（79题）.docx',
 'X2': '人教B版选必1 第2章 平面解析几何·衔接件（13题）.docx',
 'I2': '人教B版选必1 第2章 平面解析几何·知识清单（完成）.docx',
 'E': '人教B版选必1 第2章 平面解析几何（2.1—2.3.3）·讲练件（92题）.docx',
 'F': '人教B版选必1 第2章 平面解析几何（2.3.4—2.5.2）·讲练件（90题）.docx',
 'G': '人教B版选必1 第2章 平面解析几何（2.6.1—2.7.2）·讲练件（68题）.docx',
 'H': '人教B版选必1 第2章 平面解析几何（2.8）·讲练件（89题）.docx',
}
def runinfo(r):
    rpr = r._r.rPr
    b = sz = shd = None
    if rpr is not None:
        be = rpr.find(qn('w:b'))
        b = be is not None and be.get(qn('w:val')) not in ('0', 'false')
        se = rpr.find(qn('w:sz'))
        if se is not None: sz = se.get(qn('w:val'))
        she = rpr.find(qn('w:shd'))
        if she is not None: shd = she.get(qn('w:fill'))
    return b, sz, shd

def find_para(tag, prefix):
    doc = Document(os.path.join(BASE, F[tag]))
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith(prefix):
            print('[%s] 段%03d TEXT: %s' % (tag, i, p.text[:70]))
            for r in p.runs[:4]:
                b, sz, shd = runinfo(r)
                print('     RUN b=%s sz=%s shd=%s | %s' % (b, sz, shd, r.text[:40]))
            return i
    print('[%s] 未找到前缀 %s' % (tag, prefix))
    return None

print('== 抽测3件（A4点名） ==')
find_para('H', '2.8.2-1．')
find_para('H', '2.8-1．')
# X1首个题号段
doc = Document(os.path.join(BASE, F['X1']))
for i, p in enumerate(doc.paragraphs):
    if re.match(r'^\d+\.\d+\.\d+\.\d+-\d+．', p.text):
        print('[X1] 首个题号段%03d TEXT: %s' % (i, p.text[:70]))
        for r in p.runs[:4]:
            b, sz, shd = runinfo(r)
            print('     RUN b=%s sz=%s shd=%s | %s' % (b, sz, shd, r.text[:40]))
        break

print('== 全6讲练件 C9C9C9题号run全量扫描 ==')
pat = re.compile(r'^\d+\.\d+(\.\d+)*-\d+．$')
res = {}
for tag in ['B', 'C', 'E', 'F', 'G', 'H']:
    doc = Document(os.path.join(BASE, F[tag]))
    hits = []
    for i, p in enumerate(doc.paragraphs):
        for r in p.runs:
            b, sz, shd = runinfo(r)
            if shd == 'C9C9C9' and pat.match(r.text.strip()):
                hits.append({'para': i, 'text': r.text, 'b': b, 'sz': sz})
    res[tag] = hits
    print('[%s] C9C9C9题号run = %d 例: %s' % (tag, len(hits), [h['text'] for h in hits[:6]]))
    for h in hits[:3]:
        print('     段%03d b=%s sz=%s | %s' % (h['para'], h['b'], h['sz'], h['text']))

print('== 清单件条目号/基进计数 ==')
res2 = {}
for tag in ['I1', 'I2']:
    doc = Document(os.path.join(BASE, F[tag]))
    n条目 = n基 = n进 = 0
    pate = re.compile(r'^\d+\.\d+-\d+．$')
    for p in doc.paragraphs:
        for r in p.runs:
            b, sz, shd = runinfo(r)
            if shd == 'C9C9C9' and pate.match(r.text.strip()):
                n条目 += 1
        t = p.text
        if '〔基〕' in t: n基 += 1
        if '〔进〕' in t: n进 += 1
    res2[tag] = {'条目号灰底run': n条目, '基': n基, '进': n进}
    print('[%s] 条目号灰底run=%d 〔基〕=%d 〔进〕=%d' % (tag, n条目, n基, n进))

json.dump({'讲练件C9C9C9题号': {k: v for k, v in res.items()}, '清单计数': res2},
          open(r'C:\提示词\工作区\同步-数学选必1复合修复-0903\子步8\tmp\探针_收尾.json', 'w', encoding='utf-8'), ensure_ascii=False, indent=1)
print('探针_收尾.json 落盘')
