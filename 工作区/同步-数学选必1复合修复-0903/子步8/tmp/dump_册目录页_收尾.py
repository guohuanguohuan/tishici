# -*- coding: utf-8 -*-
"""册目录页逐行核查：dump全部段落(python-docx口径)＋行型谓词判定＋三源恒等核验"""
import sys, io, json, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
from docx import Document
from docx.oxml.ns import qn

PATH = r'C:\提示词\高中数学\高中数学同步\人教B版选必1·册目录页.docx'
doc = Document(PATH)

def pprops(p):
    pPr = p._p.pPr
    ind = shd = tabs = jc = None
    if pPr is not None:
        i = pPr.find(qn('w:ind'))
        if i is not None:
            ind = {k.split('}')[1]: v for k, v in i.attrib.items()}
        s = pPr.find(qn('w:shd'))
        if s is not None:
            shd = s.get(qn('w:fill'))
        t = pPr.find(qn('w:tabs'))
        if t is not None:
            tabs = [{k.split('}')[1]: v for k, v in tb.attrib.items()} for tb in t.findall(qn('w:tab'))]
        j = pPr.find(qn('w:jc'))
        if j is not None:
            jc = j.get(qn('w:val'))
    return ind, shd, tabs, jc

def classify(text, ind, shd, runs_bold):
    # 行型谓词（按序首中即定）：章行/件级行/节级行/小注行；空行与未归类单列
    if re.search(r'第\d+章', text) and runs_bold and 'P' not in text and '·本' not in text and '页' not in text.replace('第%d章',''):
        pass
    if re.search(r'第\d+章', text) and runs_bold and not re.search(r'·本\d+页', text) and not re.search(r'P\d+', text):
        return '章行'
    if ind and ind.get('left') == '420' and shd == 'C9C9C9' and re.search(r'·本\d+页', text):
        return '件级行'
    if re.match(r'^\d+\.\d+', text) and ind and ind.get('left') == '840':
        return '节级行'
    if '页码＝所在本部分内页码' in text:
        return '小注行'
    if text.strip() == '':
        return '空行'
    return '未归类'

rows = []
for i, p in enumerate(doc.paragraphs):
    text = p.text
    ind, shd, tabs, jc = pprops(p)
    bolds = []
    for r in p.runs:
        rpr = r._r.rPr
        b = rpr is not None and rpr.find(qn('w:b')) is not None and (rpr.find(qn('w:b')).get(qn('w:val')) not in ('0', 'false'))
        sz = None
        if rpr is not None:
            szel = rpr.find(qn('w:sz'))
            if szel is not None:
                sz = szel.get(qn('w:val'))
        rshd = None
        if rpr is not None:
            sel = rpr.find(qn('w:shd'))
            if sel is not None:
                rshd = sel.get(qn('w:fill'))
        bolds.append({'t': r.text, 'b': b, 'sz': sz, 'shd': rshd})
    allbold = bool(p.runs) and all(x['b'] for x in bolds if x['t'].strip())
    typ = classify(text, ind, shd, allbold)
    rows.append({'i': i, 'type': typ, 'text': text, 'ind': ind, 'pshd': shd,
                 'tabs': tabs, 'jc': jc, 'runs': bolds})

cnt = {}
for r in rows:
    cnt[r['type']] = cnt.get(r['type'], 0) + 1
print('段落总数(python-docx document.paragraphs) =', len(rows))
print('行型计数 =', json.dumps(cnt, ensure_ascii=False))
for r in rows:
    print('---[%02d] %s | ind=%s pshd=%s jc=%s tabs=%s' % (r['i'], r['type'], r['ind'], r['pshd'], r['jc'], r['tabs']))
    print('    TEXT: %s' % r['text'])
    for x in r['runs']:
        print('    RUN b=%s sz=%s shd=%s | %s' % (x['b'], x['sz'], x['shd'], x['t'][:50]))

json.dump(rows, open(r'C:\提示词\工作区\同步-数学选必1复合修复-0903\子步8\tmp\dump_册目录页_收尾.json', 'w', encoding='utf-8'), ensure_ascii=False, indent=1)

# 节页面设置（版式）：A4/边距/分栏/制表位停靠
sec = doc.sections[0]
print('=== 版式: pgSz=%sx%s pgMar L%s R%s T%s B%s | cols=%s' % (
    sec.page_width, sec.page_height, sec.left_margin, sec.right_margin,
    sec.top_margin, sec.bottom_margin,
    sec._sectPr.find(qn('w:cols')).get(qn('w:num')) if sec._sectPr.find(qn('w:cols')) is not None else None))
