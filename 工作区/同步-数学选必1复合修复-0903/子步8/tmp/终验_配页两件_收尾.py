# -*- coding: utf-8 -*-
"""终态验证：两配页件最终dump＋T2图例同构断言（run级对实物）＋自检八项XML实测数字"""
import sys, io, os, json, re, zipfile
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
from docx import Document
from docx.oxml.ns import qn

BASE = r'C:\提示词\高中数学\高中数学同步'
P1 = os.path.join(BASE, '人教B版选必1·册目录页.docx')
P2 = os.path.join(BASE, '人教B版选必1·使用说明.docx')

def runparams(r):
    rpr = r._r.rPr
    b = sz = shd = None
    if rpr is not None:
        be = rpr.find(qn('w:b'))
        b = be is not None and be.get(qn('w:val')) not in ('0', 'false')
        se = rpr.find(qn('w:sz'))
        if se is not None: sz = se.get(qn('w:val'))
        she = rpr.find(qn('w:shd'))
        if she is not None: shd = she.get(qn('w:fill'))
    return b, sz, shd

print('===== T2 图例同构断言（使用说明 ↔ 实物） =====')
doc2 = Document(P2)
paras = doc2.paragraphs
checks = []
def para_by(prefix):
    for i, p in enumerate(paras):
        if p.text.startswith(prefix): return i, p
    return None, None

# 实物基准值（探针_收尾.json + 实物run探针_子步8.json 实测）
probe = json.load(open(r'C:\提示词\工作区\同步-数学选必1复合修复-0903\子步8\tmp\探针_收尾.json', encoding='utf-8'))

# ①行4 题型级题号块 ↔ H 2.8.2-1（白底加粗）
i, p = para_by('1.1.1.1-1．')
b, sz, shd = runparams(p.runs[0])
checks.append(('行4题型级题号块', (b, sz, shd), (True, '24', None), '实物H段160: b=True sz=24 shd=None'))
# ②行5 衔接件题号块 ↔ X1 1.2.1.1-1（C9C9C9加粗）
i, p = para_by('1.2.1.1-1．')
b, sz, shd = runparams(p.runs[0])
checks.append(('行5衔接件题号块', (b, sz, shd), (True, '24', 'C9C9C9'), '实物X1段033: b=True sz=24 shd=C9C9C9'))
# ③【答案】芯片行 ↔ B/H 芯片（黑字白底不加粗）
i, p = para_by('【答案】')
b, sz, shd = runparams(p.runs[0])
checks.append(('芯片【答案】行', (b, sz, shd), (False, '24', None), '实物B/H芯片: b=False sz=24 shd=None'))
# ④讲部需背灰底行 ↔ H讲部灰底run
i, p = para_by('定义：在空间')
gr = [runparams(r) for r in p.runs if runparams(r)[2] == 'C9C9C9']
checks.append(('讲部灰底行(大小/方向)', gr[:2], [(False, '24', 'C9C9C9')] * len(gr), '实物H讲部: b=False sz=24 shd=C9C9C9'))
# ⑤条目号行 ↔ I2条目号
i, p = para_by('2.4-13．')
b, sz, shd = runparams(p.runs[0])
checks.append(('条目号2.4-13．', (b, sz, shd), (False, '24', 'C9C9C9'), '实物I2条目号: C9C9C9 b=False sz=24'))
# ⑥第一子层（1）
i, p = para_by('（1）定理内容')
b, sz, shd = runparams(p.runs[0])
checks.append(('第一子层（1）', (b, sz, shd), (False, '24', 'C9C9C9'), '实物B（1）: C9C9C9 b=False sz=24'))
# ⑦题干底纹段
i, p = para_by('如图，在正方体')
pshd = p._p.pPr.find(qn('w:shd')).get(qn('w:fill'))
checks.append(('题干底纹段', pshd, 'E0E0E0', '§7题干底纹E0E0E0'))
# ⑧定理细框段
i, p = para_by('定理：三条射线')
pbdr = p._p.pPr.find(qn('w:pBdr')) is not None
pshd2 = p._p.pPr.find(qn('w:shd')).get(qn('w:fill'))
checks.append(('定理细框段(pBdr+灰底)', (pbdr, pshd2), (True, 'C9C9C9'), '§7灰底＋细框双标记'))

allok = True
for name, got, exp, src in checks:
    ok = got == exp
    allok &= ok
    print('%s %s: 实测=%s 期望=%s 源=%s' % ('✓' if ok else '✗', name, got, exp, src))
print('同构断言总判:', 'PASS' if allok else 'FAIL')

print('===== 使用说明 §11三要素 =====')
txts = [p.text for p in paras]
def has(s): return any(s in t for t in txts)
for label, cond in [
    ('①视觉锚图例区(一、)', has('一、视觉锚图例')),
    ('〔基〕图例句', any(t.startswith('〔基〕＝基础必会') for t in txts)),
    ('〔进〕图例句', any(t.startswith('〔进〕＝进阶汇总') for t in txts)),
    ('②难度三档(二、)', has('二、难度三档与提分线')),
    ('提分线三形态', has('保60%') and has('保80%') and has('冲100%')),
    ('建议学习路径', has('建议学习路径')),
    ('答案用法小节', has('答案用法：每题先独立读题动手')),
    ('③件型用法(三、)', has('三、本册件型用法')),
    ('衔接件必会说明', has('衔接件：初升高铺垫，不区分难度、全部必会')),
    ('部分封面用法', has('部分封面：每部分')),
    ('装订组合方案A/B/C', has('方案A整册装订') and has('方案B按件型抽订') and has('方案C分层抽订')),
]:
    print('%s %s' % ('✓' if cond else '✗', label))

print('===== 两配页件 自检XML实测 =====')
for tag, fp in (('册目录页', P1), ('使用说明', P2)):
    doc = Document(fp)
    xml = zipfile.ZipFile(fp).read('word/document.xml').decode('utf-8')
    sects = doc.sections
    hf = [n for n in zipfile.ZipFile(fp).namelist() if re.search(r'(header|footer)\d*\.xml$', n)]
    sec = sects[0]
    cols = sec._sectPr.find(qn('w:cols'))
    jc_bad = 0
    for p in doc.paragraphs:
        pPr = p._p.pPr
        j = pPr.find(qn('w:jc')) if pPr is not None else None
        if j is not None and j.get(qn('w:val')) != 'left':
            jc_bad += 1
    m = {
        '段落数': len(doc.paragraphs),
        '节数': len(sects),
        '页眉页脚部件数': len(hf),
        'w:ins': xml.count('<w:ins '), 'w:del': xml.count('<w:del '),
        'w:strike': xml.count('<w:strike'), 'w:highlight': xml.count('<w:highlight'),
        '非auto色w:color': len(re.findall(r'<w:color w:val="(?!auto")', xml)),
        'pageBreakBefore': xml.count('<w:pageBreakBefore'),
        'keepNext': xml.count('<w:keepNext'), 'keepLines': xml.count('<w:keepLines'),
        '手动分页符w:br_page': len(re.findall(r'<w:br w:type="page"', xml)),
        'pgSz': '%sx%s' % (sec.page_width, sec.page_height),
        'pgMar_LR_T_B': '%s/%s/%s/%s' % (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin),
        'cols': cols.get(qn('w:num')) if cols is not None else '默认(1)',
        'jc非left段数': jc_bad,
        'titlePg': sec._sectPr.find(qn('w:titlePg')) is not None,
    }
    print(tag, json.dumps(m, ensure_ascii=False))
    # ⑧字体抽查：非数学run字体
    fonts = set()
    for rm in re.finditer(r'<w:rPr>.*?</w:rPr>', xml, re.S):
        rf = re.search(r'<w:rFonts ([^/]*)/>', rm.group(0))
        if rf: fonts.add(rf.group(1))
    print(tag, 'rFonts枚举:', list(fonts)[:4])
