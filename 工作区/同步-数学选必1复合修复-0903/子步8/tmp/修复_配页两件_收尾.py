# -*- coding: utf-8 -*-
"""收尾修复：①册目录页6件级行shd属性fill→w:fill（命名空间bug修底纹不渲染）②使用说明行4失实句据实改写（A5 else支＋题型级限定）③使用说明行35字号18→24半点"""
import sys, io, shutil, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
from docx import Document
from docx.oxml.ns import qn

BASE = r'C:\提示词\高中数学\高中数学同步'
BAK = r'C:\提示词\工作区\同步-数学选必1复合修复-0903\子步8\tmp\bak_收尾'
os.makedirs(BAK, exist_ok=True)
P1 = os.path.join(BASE, '人教B版选必1·册目录页.docx')
P2 = os.path.join(BASE, '人教B版选必1·使用说明.docx')
if not os.path.exists(os.path.join(BAK, '人教B版选必1·册目录页.docx')):
    shutil.copy2(P1, os.path.join(BAK, '人教B版选必1·册目录页.docx'))
    shutil.copy2(P2, os.path.join(BAK, '人教B版选必1·使用说明.docx'))
    print('备份落盘 bak_收尾\\')

# ① 册目录页：件级行 shd 无命名空间 fill 属性 → w:fill
doc = Document(P1)
nfix = 0
for p in doc.paragraphs:
    pPr = p._p.pPr
    if pPr is None: continue
    ind = pPr.find(qn('w:ind'))
    if ind is not None and ind.get(qn('w:left')) == '420' and re.search(r'·本\d+页', p.text):
        shd = pPr.find(qn('w:shd'))
        assert shd is not None, '件级行无shd元素: %s' % p.text[:30]
        bad = shd.get('fill')
        if shd.get(qn('w:fill')) is None and bad:
            del shd.attrib['fill']
            shd.set(qn('w:fill'), bad)
            nfix += 1
            print('  修: %s fill=%s→w:fill' % (p.text[:36], bad))
        assert shd.get(qn('w:val')) == 'clear' and shd.get(qn('w:color')) == 'auto'
assert nfix == 6, '修件数≠6: %d' % nfix
doc.save(P1)
print('册目录页：6件级行 shd w:fill=C9C9C9 已修复（整行底纹生效）')

# ② 使用说明行4失实句改写
OLD = '题号块底纹式样仅衔接件·清单件仍现行（下行即衔接件式样）。'
NEW = '题型级题号块底纹仅衔接件·清单件仍现行（下行即衔接件式样：C9C9C9＋整块加粗）；讲练件讲部填空块题号（节号-序号型，如2.8-1．）随填空块挂C9C9C9、不加粗，与题型级题号块有别。'
doc = Document(P2)
p4 = doc.paragraphs[4]
assert '题号块底纹式样仅衔接件' in p4.text, '行4锚失'
done = False
for r in p4.runs:
    if OLD in r.text:
        r.text = r.text.replace(OLD, NEW)
        done = True
assert done, '行4尾句未定位到run'
print('使用说明行4尾句：\n  前: %s\n  后: %s' % (OLD, NEW))
# ③ 行35字号18→24
p35 = doc.paragraphs[35]
assert p35.text.startswith('本册页眉页脚同串'), '行35锚失: %s' % p35.text[:20]
n35 = 0
for r in p35.runs:
    rpr = r._r.rPr
    if rpr is not None:
        se = rpr.find(qn('w:sz'))
        if se is not None and se.get(qn('w:val')) == '18':
            se.set(qn('w:val'), '24'); n35 += 1
        se2 = rpr.find(qn('w:szCs'))
        if se2 is not None and se2.get(qn('w:val')) == '18':
            se2.set(qn('w:val'), '24')
assert n35 == 1, '行35 sz=18 run数≠1: %d' % n35
doc.save(P2)
print('使用说明：行35字号18→24半点（1 run）')
