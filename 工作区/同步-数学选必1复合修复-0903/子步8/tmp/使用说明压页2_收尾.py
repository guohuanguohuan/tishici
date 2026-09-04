# -*- coding: utf-8 -*-
"""使用说明压页第二步：删空段22/27 → 测 → 精简行13注文 → 测"""
import sys, io, os, json, time, subprocess
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
from docx import Document

TMP = r'C:\提示词\工作区\同步-数学选必1复合修复-0903\子步8\tmp'
P2 = r'C:\提示词\高中数学\高中数学同步\人教B版选必1·使用说明.docx'
WORKER = os.path.join(TMP, 'com_worker_收尾.py')

def measure():
    status = os.path.join(TMP, 'com_status_m3.jsonl')
    if os.path.exists(status): os.remove(status)
    proc = subprocess.Popen([sys.executable, WORKER, status, P2])
    t0 = time.time()
    while time.time() - t0 < 90:
        if os.path.exists(status):
            for ln in open(status, encoding='utf-8'):
                if '"open_done"' in ln:
                    e = json.loads(ln); proc.wait(timeout=10); return e['pages']
        if proc.poll() is not None:
            break
        time.sleep(0.5)
    proc.kill(); return None

# 删剩余空段（文本为空的段）
doc = Document(P2)
removed = []
for p in list(doc.paragraphs):
    if p.text == '':
        removed.append(len(removed))
        p._p.getparent().remove(p._p)
doc.save(P2)
print('删空段 %d 个' % len(removed))
pg = measure()
print('删空段后 COM页数 =', pg)

if pg == 2:
    # 精简行13注文（块标签芯片说明，9pt，原文>100字折2行→压到1行）
    doc = Document(P2)
    p13 = doc.paragraphs[11]  # 删2空段后原13→11
    print('段11锚文本: %s' % p13.text[:30])
    OLD13 = None
    for r in p13.runs:
        t = r.text
        if '块标签芯片' in t:
            OLD13 = t
    assert OLD13, '行13注文未锚到'
    NEW13 = '　块标签芯片：【答案】【知识点】【分析】【详解】【点睛】【编注】等行内标签——上行＝讲练件现行形态（黑字白底·底纹已废止）；芯片底纹与答案值灰底仅衔接件·清单件仍现行。'
    if len(OLD13) <= len(NEW13):
        print('注: 原文已够短，跳过')
    else:
        for r in p13.runs:
            if r.text == OLD13:
                r.text = NEW13
        doc.save(P2)
        print('行13注文 %d字→%d字' % (len(OLD13), len(NEW13)))
    pg = measure()
    print('精简行13后 COM页数 =', pg)
print('最终页数 =', pg)
