# -*- coding: utf-8 -*-
"""使用说明压回1页：行4尾句精简（保题型级限定＋据实句）→COM测页→仍2页则删空段01→再测→仍2页删空段23→再测"""
import sys, io, os, json, time, subprocess
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
from docx import Document
from docx.oxml.ns import qn

TMP = r'C:\提示词\工作区\同步-数学选必1复合修复-0903\子步8\tmp'
P2 = r'C:\提示词\高中数学\高中数学同步\人教B版选必1·使用说明.docx'
WORKER = os.path.join(TMP, 'com_worker_收尾.py')

def measure():
    status = os.path.join(TMP, 'com_status_measure.jsonl')
    if os.path.exists(status): os.remove(status)
    proc = subprocess.Popen([sys.executable, WORKER, status, P2])
    t0 = time.time()
    while time.time() - t0 < 90:
        if proc.poll() is not None: break
        if os.path.exists(status):
            for ln in open(status, encoding='utf-8'):
                if '"open_done"' in ln:
                        e = json.loads(ln)
                        proc.wait(timeout=10)
                        return e['pages']
        time.sleep(0.5)
    proc.kill()
    return None

CUR = '题型级题号块底纹仅衔接件·清单件仍现行（下行即衔接件式样：C9C9C9＋整块加粗）；讲练件讲部填空块题号（节号-序号型，如2.8-1．）随填空块挂C9C9C9、不加粗，与题型级题号块有别。'
TRIM = '题型级题号块底纹仅衔接件·清单件仍现行（下行即衔接件式样）；讲练件讲部填空块题号（2.8-1型）随块挂C9C9C9、不加粗。'

doc = Document(P2)
p4 = doc.paragraphs[4]
done = False
for r in p4.runs:
    if CUR in r.text:
        r.text = r.text.replace(CUR, TRIM)
        done = True
assert done, '行4现文未锚到'
doc.save(P2)
print('行4尾句精简为: %s' % TRIM)
print('COM页数 =', measure())

for blank_txt_cond in (lambda t: t == '',):
    pass

def drop_blank(pred_desc):
    doc = Document(P2)
    for i, p in enumerate(doc.paragraphs):
        if p.text == '' and ((pred_desc == 'title-next' and i == 1) or (pred_desc == 'sec2-prev' and i == 23)):
            p._p.getparent().remove(p._p)
            doc.save(P2)
            print('删空段 段%d（%s）' % (i, pred_desc))
            return True
    return False

pg = measure()
print('精简后 COM页数 =', pg)
if pg == 2:
    drop_blank('title-next')
    pg = measure()
    print('删段01后 COM页数 =', pg)
if pg == 2:
    drop_blank('sec2-prev')
    pg = measure()
    print('删段23后 COM页数 =', pg)
print('最终页数 =', pg)
