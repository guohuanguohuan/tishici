# -*- coding: utf-8 -*-
"""A''册目录页重造：列头＋章行2＋件型行6（C9C9C9整行底纹＋·本N＋题量/三档）＋节级行25（缩进840＋题量括注＋部分内页码列）"""
import json, sys, os, re, zipfile
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
def q(t): return '{%s}%s' % (W, t)
BASE = r'C:\提示词\高中数学\高中数学同步'
BOOK = '人教B版选必1'

SEC2NAME = {
 '1.1': '空间向量及其运算', '1.2': '空间向量在立体几何中的应用',
 '2.1': '坐标法', '2.2': '直线', '2.3': '圆', '2.4': '曲线与方程',
 '2.5': '椭圆', '2.6': '双曲线', '2.7': '抛物线', '2.8': '直线与圆锥曲线的位置关系',
}
PARTS = [  # (部分tag, 件名, 本n, 件型行文案（含题量）, 件级页码)
 ('第1章·衔接', '人教B版选必1 第1章 空间向量与立体几何·衔接件（29题）.docx', 1,
  '衔接件（29题·全部必会）', 1),
 ('第1章·清单', '人教B版选必1 第1章 空间向量与立体几何·知识清单（完成）.docx', 2,
  '知识清单（47条：基33·进14）', 1),
 ('第1章·讲练', '人教B版选必1 第1章 空间向量与立体几何（上）·讲练件（61题）.docx', 3,
  '讲练件（140题：简单21·中档104·难15）', 1),
 ('第2章·衔接', '人教B版选必1 第2章 平面解析几何·衔接件（13题）.docx', 4,
  '衔接件（13题·全部必会）', 1),
 ('第2章·清单', '人教B版选必1 第2章 平面解析几何·知识清单（完成）.docx', 5,
  '知识清单（67条：基38·进29）', 1),
 ('第2章·讲练', '人教B版选必1 第2章 平面解析几何（2.1—2.3.3）·讲练件（92题）.docx', 6,
  '讲练件（339题：简单47·中档246·难46）', 1),
]
SEC_TITLE_RE = re.compile(r'^(\d+\.\d+(?:\.\d+)?)\s+(.+?)　本节(\d+)题')

def ptext(el): return ''.join(t.text or '' for t in el.iter(q('t')))

def sec_rows(fname, skip_before=0):
    """读讲练件节标题行（含统计段）→ [(节号, 节名, 题量)]（跳过导航表内行：表格内段落排除）。"""
    doc = etree.fromstring(zipfile.ZipFile(os.path.join(BASE, fname)).read('word/document.xml'))
    rows = []
    for p in doc.iter(q('p')):
        anc = p.getparent()
        in_tbl = False
        while anc is not None:
            if anc.tag == q('tbl'): in_tbl = True; break
            anc = anc.getparent()
        if in_tbl: continue
        m = SEC_TITLE_RE.match(ptext(p).strip())
        if m and m.group(1) not in [r[0] for r in rows]:
            rows.append((m.group(1), m.group(2), int(m.group(3))))
    return rows

def add_para(doc, text_runs, sz=18, bold=False, shd=None, ind=None, page=None, sec=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = Pt(14)
    # 页码列：点线制表位右停靠
    pf.tab_stops.add_tab_stop(Cm(17.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if shd:
        pPr = p._p.get_or_add_pPr()
        shd_el = OxmlElement('w:shd')
        shd_el.set(qn('w:val'), 'clear'); shd_el.set(qn('w:color'), 'auto'); shd_el.set(qn('w:fill'), shd)
        pPr.append(shd_el)
    if ind is not None:
        pPr = p._p.get_or_add_pPr()
        ind_el = OxmlElement('w:ind'); ind_el.set(qn('w:left'), str(ind)); pPr.append(ind_el)
    if isinstance(text_runs, str):
        text_runs = [(text_runs, bold)]
    for text, b in text_runs:
        r = p.add_run(text)
        r.font.size = Pt(sz / 2)
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.bold = b
    if page is not None:
        r = p.add_run('\t%d' % page)
        r.font.size = Pt(sz / 2)
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.bold = bold
    return p

def main(out):
    loc = json.load(open('wip/节定位-纯.json', encoding='utf-8'))
    sec2page = {}
    for f in loc['files']:
        for s_ in f.get('sections', []):
            sec2 = s_['no'].rsplit('.', 1)[0] if s_['no'].count('.') > 1 else s_['no']
            if sec2 not in sec2page or s_['part_page'] < sec2page[sec2]:
                sec2page[sec2] = s_['part_page']
    doc = docx.Document()
    add_para(doc, [(BOOK + '·册目录页', True)], sz=32)
    add_para(doc, '件／节（括注＝题量）', sz=18)
    # —— 第1章 ——
    add_para(doc, [('第1章 空间向量与立体几何', True)], sz=24, page=1)
    for tag, fn, bn, label, pg in PARTS[:3]:
        add_para(doc, [(label, False), ('　·本%d' % bn, False)], sz=18, shd='C9C9C9', ind=420, page=pg)
    agg1 = agg_sec2([PARTS[2][1],
                     '人教B版选必1 第1章 空间向量与立体几何（下）·讲练件（79题）.docx'])
    for sec2, n in agg1:
        add_para(doc, [('%s %s（%d题）' % (sec2, SEC2NAME.get(sec2, sec2), n), False)],
                 sz=18, ind=840, page=sec2page.get(sec2, 1))
    # —— 第2章 ——
    add_para(doc, [('第2章 平面解析几何', True)], sz=24, page=1)
    for tag, fn, bn, label, pg in PARTS[3:]:
        add_para(doc, [(label, False), ('　·本%d' % bn, False)], sz=18, shd='C9C9C9', ind=420, page=pg)
    ch2 = ['人教B版选必1 第2章 平面解析几何（2.1—2.3.3）·讲练件（92题）.docx',
           '人教B版选必1 第2章 平面解析几何（2.3.4—2.5.2）·讲练件（90题）.docx',
           '人教B版选必1 第2章 平面解析几何（2.6.1—2.7.2）·讲练件（68题）.docx',
           '人教B版选必1 第2章 平面解析几何（2.8）·讲练件（89题）.docx']
    agg2 = agg_sec2(ch2)
    for sec2, n in agg2:
        add_para(doc, [('%s %s（%d题）' % (sec2, SEC2NAME.get(sec2, sec2), n), False)],
                 sz=18, ind=840, page=sec2page.get(sec2, 1))
    sect = doc.sections[0]
    sp_el = sect._sectPr
    for tag, attrs in (('w:pgSz', {'w:w': '11906', 'w:h': '16838'}),
                       ('w:pgMar', {'w:top': '850', 'w:right': '850', 'w:bottom': '850', 'w:left': '850',
                                    'w:header': '850', 'w:footer': '850', 'w:gutter': '0'})):
        el = sp_el.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); sp_el.append(el)
        for k, v in attrs.items():
            el.set(qn(k), v)
    doc.core_properties.title = BOOK + '·册目录页'
    doc.core_properties.author = ''
    doc.save(out)
    print('saved', out)


def agg_sec2(files):
    """多卷三级节题量聚合到二级：[(sec2, N)] 按序。"""
    agg = {}
    for fn in files:
        for sec, name, nq in sec_rows(fn):
            sec2 = sec.rsplit('.', 1)[0] if sec.count('.') > 1 else sec
            agg[sec2] = agg.get(sec2, 0) + nq
    return sorted(agg.items(), key=lambda x: [int(t) for t in x[0].split('.')])


if __name__ == '__main__':
    main(sys.argv[1] if len(sys.argv) > 1 else 'wip/册目录页.docx')
