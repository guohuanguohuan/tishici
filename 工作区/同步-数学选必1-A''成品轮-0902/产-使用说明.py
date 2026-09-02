# -*- coding: utf-8 -*-
"""A''使用说明页重做：§11三要素（视觉锚图例A''口径／难度三档＋答案用法／件型用法＋方案A—C）"""
import sys, os
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BOOK = '人教B版选必1'
FILL = 'C9C9C9'; T1 = 'ADC2DA'; T2 = 'C6D4E3'; STEM = 'E0E0E0'
SZ_TITLE = 32; SZ_BODY = 24; SZ_SMALL = 18

def para(doc, runs, sz=SZ_BODY, shade=None, bold=False, border=False):
    """runs: [(text, {shd, bold, color})] 或 str"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = Pt(20.5)
    if isinstance(runs, str):
        runs = [(runs, {})]
    if shade:                       # 段级底纹（题干/标题式样行）
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), shade)
        pPr.append(shd)
    if border:                      # 定理细框
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement('w:' + side)
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:space'), '4'); b.set(qn('w:color'), 'auto')
            pBdr.append(b)
        pPr.append(pBdr)
    for text, sty in runs:
        r = p.add_run(text)
        r.font.size = Pt(sz / 2)
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.bold = bool(sty.get('bold', bold))
        if sty.get('shd'):
            from docx.oxml import OxmlElement
            rPr = r._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), sty['shd'])
            rPr.append(shd)
    return p

def main(out):
    doc = docx.Document()
    para(doc, [(BOOK + '·使用说明', {'bold': True})], sz=SZ_TITLE)
    para(doc, '')
    para(doc, [('一、视觉锚图例', {'bold': True})], sz=SZ_BODY, shade=T1)
    para(doc, '（式样＝本册现行版式真实呈现；正文双栏排版——每栏约8.6cm、栏间有分隔线，卷首标题与导航表通栏单栏。）', sz=SZ_SMALL)
    # 题号块
    para(doc, [('1.1.1.1-1．', {'shd': FILL, 'bold': True}), ('（简单·保60%·卡壳看答案）', {'bold': True}),
               ('　题号块：「题型号-节内序号．（档位·提分线·卡壳看答案）」——底纹盖整个题号、括注不挂底纹、整块加粗；序号＝节内连续（同节跨题型累进）；题间不留空行，靠底纹分隔。', {})])
    para(doc, [('1.2.1.1-1．', {'shd': FILL, 'bold': True}), ('（衔接必会·卡壳看答案）', {'bold': True}),
               ('　衔接件不区分难度——两段式括注、全部必会。', {})])
    # 标题
    para(doc, '1.1 空间向量及其运算', sz=28, shade=T1, bold=True)
    para(doc, '　章标题与节标题：整行铺底#ADC2DA＋加粗＋顶格（章三号／节四号），章标题另加底边框通栏细线。', sz=SZ_SMALL)
    para(doc, [('1.1.1 空间向量的概念：向量的表示', {'bold': True})], sz=SZ_BODY, shade=T2)
    para(doc, [('1.1.1.1 空间向量的概念：向量加减法的三角形法则', {'bold': True}),
               ('　2题：1.1.1.1-1～1.1.1.1-2', {})], sz=SZ_BODY, shade=T2)
    para(doc, '　讲部标题与题型标题：整行铺底#C6D4E3＋加粗＋顶格（均小四）；题型标题行末带统计段「　N题：题号a～题号b」。', sz=SZ_SMALL)
    # 题干底纹
    para(doc, '如图，在正方体ABCD-A₁B₁C₁D₁中，点E是AB的中点。（题干整体铺底示例）', shade=STEM)
    para(doc, '　题干底纹：每题题号块起至【答案】行之前的全部段落铺#E0E0E0（含选项与设问）；配图独立段不铺；解析区白底。', sz=SZ_SMALL)
    # 答案值＋块标签
    para(doc, [('【答案】', {'shd': FILL}), (' ', {}), ('C', {'shd': FILL}), ('　', {}),
               ('【知识点】', {'shd': FILL}), (' ', {}), ('1.1.1 空间向量的概念', {})])
    para(doc, '　块标签芯片：【答案】【知识点】【分析】【详解】【点睛】【编注】【大招指引】【题后反思】等行内栏目标签挂同款灰底；其后答案值与需背内容一律挂同款灰底（公式型随公式整体挂灰）。', sz=SZ_SMALL)
    # 条目号＋子层＋基进
    para(doc, [('2.4-13．', {'shd': FILL}), ('〔基〕', {}), ('平面向量基本定理', {})])
    para(doc, [('（1）', {'shd': FILL}), ('定理内容：不共线两向量可线性表出任一向量。', {})])
    para(doc, '　条目号「节号-序号．」与条目第一子层「（N）」挂灰块；〔基〕＝基础必会、〔进〕＝进阶汇总（不挂底纹）。', sz=SZ_SMALL)
    para(doc, '〔基〕＝基础必会：必须学完本条目，才能做本章题目')
    para(doc, '〔进〕＝进阶汇总：本章各题型常识/结论的汇总，方便复习，必须先做题再回看')
    # 定理框
    para(doc, '定理：三条射线两两垂直且相等时，PA⊥平面ABC（三垂线定理示例——定理级条目「灰底＋细框」双标记）', shade=FILL, border=True)
    para(doc, '')
    para(doc, [('二、难度三档与提分线', {'bold': True})], sz=SZ_BODY, shade=T1)
    para(doc, '简单＝单一知识点直接套用、一两步完成——保60%；中档＝常规解法多步转化或跨一两个知识点——保80%；难＝特别难的题（综合、参数、压轴）——冲100%。简单＋中档合计具备本章高考约80%分数目标所需能力。')
    para(doc, '建议学习路径：先做衔接件（必会）→过知识清单〔基〕→讲练件由易到难逐题型推进→错题隔天重做。')
    para(doc, [('答案用法：', {'bold': True}), ('每题先独立读题动手。卡壳超过10分钟立即看答案学方法，不死磕。看懂后遮住答案重做一遍，隔天重做仍错的题记入', {}),
               ('自己的错题本', {'shd': FILL}), ('。', {})])
    para(doc, '')
    para(doc, [('三、本册件型用法', {'bold': True})], sz=SZ_BODY, shade=T1)
    para(doc, '衔接件：初升高铺垫，不区分难度、全部必会，题号块为两段式「（衔接必会·卡壳看答案）」。')
    para(doc, '知识清单：〔基〕条目学完才能做题；〔进〕条目是本章各题型常识/结论汇总，必须先做题再回看。')
    para(doc, '讲练件：按教材节序由易到难装配，题型通式句在该组首题前（【编注】起段）；讲部「方法讲解｜主题」置于其题型组之前。')
    para(doc, '部分封面：每部分（同一章同一种产出的全部文件）开头一页，标注该章全部部分与本部分统计——翻书定位用。')
    para(doc, [('装订组合方案：', {'bold': True}), ('方案A整册装订（配页三件＋各部分各带部分封面）；方案B按件型抽订（全部衔接/全部清单/全部讲练）；方案C分层抽订（分层卷派生后：简单卷/中档卷/冲刺卷各按档抽订成套——分层卷以样册通过后另行派生发放）。', {})])
    para(doc, '本册页眉页脚同串：「羿郭工作室·册名 第X章 章名·件型（共N页）·本n/共M本　节号节名　第X页」——件标识随部分独立页码、本号随装订本。', sz=SZ_SMALL)
    # A4＋边距＋无页眉页脚
    sect = doc.sections[0]
    sp = sect._sectPr
    from docx.oxml import OxmlElement
    for tag, attrs in (('w:pgSz', {'w:w': '11906', 'w:h': '16838'}),
                       ('w:pgMar', {'w:top': '850', 'w:right': '850', 'w:bottom': '850', 'w:left': '850',
                                    'w:header': '850', 'w:footer': '850', 'w:gutter': '0'})):
        el = sp.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); sp.append(el)
        for k, v in attrs.items():
            el.set(qn(k), v)
    name = BOOK + '·使用说明'
    doc.core_properties.title = name
    doc.core_properties.author = ''
    doc.save(out)
    print('saved', out)

if __name__ == '__main__':
    main(sys.argv[1] if len(sys.argv) > 1 else '使用说明.docx')
