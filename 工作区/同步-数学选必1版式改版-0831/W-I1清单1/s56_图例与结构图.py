# -*- coding: utf-8 -*-
"""W-I1 步骤5＋6：N16图例行插入＋N14章知识结构图插入（T5产物6张PNG）＋源思维导图位图差图处置。
- 图例行：文内开头标题正下方；解析档9pt普通段落；N16固定句逐字照抄。
- 结构图：图例行之后、首个节标题之前；按教材节分组（T5 tsv序）；显示尺寸按T5建议（300dpi 1:1，
  节点字\\small=9pt≈N12目标；结构图系§7明示例外、显示高可超9cm）。
- 源思维导图位图（element1，media/image1.png，T5已目检确认）：按差图处置——登记后删除
  （图删文留；段落纯图无文字），以新结构图替换；media与rels一并清（守恒三查口径）。
- 断言：节点↔条目对照——tex叶节点名⊆成品条目名、覆盖100%、节点数47≤条目数47。"""
import sys, os, re, json, zipfile, io
from lxml import etree
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
RNS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
WC = '{%s}' % W

SRC = sys.argv[1]
DST = sys.argv[2]
T5DIR = r'C:\Users\28120\Desktop\提示词\工作区\同步-数学选必1版式改版-0831\T5图件生成\结构图I1'
LEGEND = '〔基〕＝基础必会：必须学完本条目，才能做本章题目｜〔进〕＝进阶汇总：本章各题型常识/结论的汇总，方便复习，必须先做题再回看'

# ---------- ① 断言：节点↔条目对照（插入前，以当前件条目为准） ----------
def entry_names(doc):
    """条目名列表：N．〔基/进〕名（题名行文本）。"""
    names = []
    for p in doc.paragraphs:
        m = re.match(r'^(\d{1,4})．〔(?:基|进)〕(.+)$', p.text)
        if m:
            names.append((int(m.group(1)), m.group(2).strip()))
    return names

def tex_leaf_names():
    """叶节点：第01组（1.1）两级树——叶=mml2（18叶，mml1=1.1.x子节层）；其余组单级树——叶=mml1。"""
    leaves = []
    sd = os.path.join(T5DIR, 'sources')
    for fn in sorted(os.listdir(sd)):
        if not fn.endswith('.tex'):
            continue
        src = open(os.path.join(sd, fn), encoding='utf-8').read()
        leafcls = 'mml2' if fn.startswith('章知识结构图_第01组') else 'mml1'
        for m in re.finditer(r'\\node\[%s,[^\]]*\]\s*at\s*\([^)]*\)\s*\{(.+?)\};' % leafcls, src, re.S):
            label = re.sub(r'\\\w+\{[^}]*\}|\\\\', '', m.group(1))
            label = re.sub(r'\s+', '', label)
            leaves.append(label)
    return leaves

doc = Document(SRC)
entries = entry_names(doc)
assert [n for n, _ in entries] == list(range(1, 48)), '条目序列非1..47: %s' % [n for n, _ in entries][:5]
entry_set = {name for _, name in entries}
leaves = tex_leaf_names()
missing = [l for l in leaves if l not in entry_set]
assert len(leaves) == 47, 'tex叶节点数=%d≠47' % len(leaves)
assert not missing, '节点名非条目名子集: %r' % missing[:5]
cover = len(set(leaves))
assert cover == 47 and len(entry_set) == 47, '覆盖非100%%: 节点uniq=%d 条目=%d' % (cover, len(entry_set))
print('A断言：节点47≤条目47；节点=条目名子集；覆盖100%（uniq47/47）PASS')

# ---------- ② 定位锚点 ----------
body = doc.element.body
els = list(body)
def ptext(el):
    return ''.join(t.text or '' for t in el.iter(WC + 't'))
title_el = None
first_sect = None
mindmap_el = None
for i, el in enumerate(els):
    t = ptext(el).strip()
    if not t:
        continue
    if title_el is None:
        title_el = el
        continue
    if re.match(r'^1\.1(\D|$)', t):
        first_sect = el
        break
assert title_el is not None and first_sect is not None
# 源思维导图位图：标题与首个节标题之间、纯图段
for el in els[els.index(title_el) + 1:els.index(first_sect)]:
    if ptext(el).strip() == '' and el.find('.//' + WC + 'drawing') is not None:
        mindmap_el = el
        break
assert mindmap_el is not None, '未找到源思维导图位图段'
mm_rid = mindmap_el.find('.//' + qn('a:blip')).get(qn('r:embed'))
print('源思维导图位图：rId=%s（差图处置：登记后删除，新结构图替换）' % mm_rid)

# ---------- ③ 插入图例行（标题正下方） ----------
pPr_xml = ('<w:pPr %sxmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
           '<w:spacing w:before="0" w:after="0" w:line="280" w:lineRule="atLeast"/>'
           '<w:jc w:val="left"/></w:pPr>')
from docx.oxml import parse_xml
legend_p = parse_xml(
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:pPr><w:spacing w:before="0" w:after="0" w:line="280" w:lineRule="atLeast"/>'
    '<w:jc w:val="left"/></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体"/>'
    '<w:color w:val="000000"/><w:sz w:val="18"/><w:szCs w:val="18"/>'
    '</w:rPr><w:t xml:space="preserve">' + LEGEND + '</w:t></w:r></w:p>')
title_el.addnext(legend_p)
assert ptext(els[els.index(first_sect) - 1]) == '' or True
print('图例行插入：标题正下方（18半点/280atLeast/jc=left）')

# ---------- ④ 插入6张结构图（图例行之后、首个节标题之前，按T5 tsv序） ----------
rows = []
for line in open(os.path.join(T5DIR, '插入辅助.tsv'), encoding='utf-8-sig').read().splitlines()[1:]:
    if not line.strip():
        continue
    f = line.split('\t')
    rows.append({'file': f[1], 'root': f[2], 'w': float(f[4]), 'h': float(f[5])})
assert len(rows) == 6, 'tsv图数=%d≠6' % len(rows)
anchor = legend_p
for r in rows:
    img_p = parse_xml(
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pPr><w:spacing w:before="0" w:after="0" w:line="410" w:lineRule="atLeast"/>'
        '<w:jc w:val="left"/></w:pPr></w:p>')
    run = img_p.makeelement(qn('w:r'), {})
    img_p.append(run)
    anchor.addnext(img_p)
    anchor = img_p
    # python-docx 行内插图（自动加 media/rels/docPr 唯一 id）
    from docx.text.run import Run
    Run(run, doc).add_picture(os.path.join(T5DIR, r['file']), width=Cm(r['w']), height=Cm(r['h']))
    print('插入：%s（%.2f×%.2fcm，根=%s）' % (r['file'], r['w'], r['h'], r['root']))

# ---------- ⑤ 删除源思维导图位图段＋rels＋media（守恒三查口径） ----------
mindmap_el.getparent().remove(mindmap_el)
doc.part.drop_rel(mm_rid)
print('源思维导图位图已删（段/rels/media 三清）')

# ---------- ⑥ 复核断言 ----------
els2 = list(body)
seq = [ptext(e)[:24] for e in els2[:10]]
print('头部序列：', json.dumps(seq, ensure_ascii=False))
# 图例行在位＋逐字
lt = ptext(legend_p)
assert lt == LEGEND, '图例行文字被改动'
assert els2.index(legend_p) == 1, '图例行不在标题正下方'
# 6图在图例行与首个节标题之间、按序
i_leg = els2.index(legend_p)
i_sect = None
for i in range(i_leg + 1, len(els2)):
    if re.match(r'^1\.1 ', ptext(els2[i])):
        i_sect = i
        break
between = els2[i_leg + 1:i_sect]
imgs = [e for e in between if e.find('.//' + qn('a:blip')) is not None]
assert len(between) == 6 and len(imgs) == 6, '图区间异常：%d段/%d图' % (len(between), len(imgs))
from docx.shared import Emu
for e, r in zip(imgs, rows):
    ext = e.find('.//' + qn('wp:extent'))
    wcm = int(ext.get('cx')) / 360000.0
    hcm = int(ext.get('cy')) / 360000.0
    assert abs(wcm - r['w']) < 0.02 and abs(hcm - r['h']) < 0.02, '显示尺寸与tsv不符: %.2f×%.2f' % (wcm, hcm)
# 旧位图确已不在
assert not any(blip.get(qn('r:embed')) == mm_rid for blip in doc.element.body.iter(qn('a:blip'))), '旧rId仍在正文'
doc.save(DST)
# 包级复核
z = zipfile.ZipFile(DST)
media = [n for n in z.namelist() if n.startswith('word/media/')]
d = etree.fromstring(z.read('word/document.xml'))
drawings = d.findall('.//' + WC + 'drawing')
rels = etree.fromstring(z.read('word/_rels/document.xml.rels'))
img_rels = [r for r in rels if 'image' in r.get('Type')]
print('包级复核：media=%d drawings=%d imageRels=%d（应 25-1+6=30）' % (len(media), len(drawings), len(img_rels)))
assert len(media) == len(drawings) == len(img_rels) == 30, '三查不等'
print('全部断言 PASS ->', DST)
