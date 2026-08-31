# -*- coding: utf-8 -*-
"""W-I1 步骤7：N23定理框——逐条目判定「定理/公式级需背条目」，是者对承载定理/公式/结论陈述的
段落加 w:pBdr 整段细框（single sz=4 space=4 auto 四边）；普通需背词不加（防满页框）。
判定与框目标（p索引＝doc.paragraphs 序，见 条目段落清单.txt）：
  YES 27条：5,7,9,10,16,17,18,22,23,26,27,28,29,30,33,34,35,36,37,38,39,41,43,44,45,46,47
  NO 20条：1,2,3,4,6,8,11,12,13,14,15,19,20,21,24,25,31,32,40,42（理由见 判定表md）
框段45个：仅定理/公式/结论陈述段；不含【编注】【微提醒】、证明/推导段、方法叙述段、（N）小标题、图段。"""
import sys, json
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WC = '{%s}' % W

SRC, DST = sys.argv[1], sys.argv[2]

# (p索引, 期望文本前缀——防错位断言)
BOXES = [
    (25, '对任意两个空间向量'), (30, '（1）定义：已知两个非零向量'), (32, '（2）由数量积的定义'),
    (44, '（2）空间向量共面的充要条件'), (48, '定理：如果三个向量'),
    (74, '设'), (77, '设'), (80, '在空间直角坐标系中，设'),
    (98, '设'), (101, '设'), (115, '设'), (118, '设'), (121, '设'), (124, '设'),
    (127, '（1）三垂线定理'), (128, '（2）三垂线定理的逆定理'),
    (142, '如图所示，O为平面α内一点'), (145, '推论：最小角定理'),
    (150, '若两异面直线'), (153, '设直线的方向向量为'), (155, '如图，若'), (158, '平面与平面相交'),
    (163, '（1）射影面积法求二面角'),
    (173, '（1）【定理】'), (177, '【常用结论】'),
    (180, '（1）【定理】'), (183, '（2）【拓展】'),
    (188, '已知直线l的单位方向向量为'), (194, '即：点A到平面'), (196, '直线与平面'), (197, '两平行平面'),
    (202, '（1）高'), (203, '（2）记相邻两个面的二面角'), (204, '（3）正四面体外接球和内切球'),
    (205, '（4）顶点在底面的射影'), (206, '（5）对棱垂直'),
    (217, '球截面性质'), (224, '设长方体的长、宽、高'), (227, '设三棱柱的高为'),
    (230, '（1）截面法'), (231, '（2）公式法'),
    (236, '设二面角'), (238, '设二面角'), (240, '设二面角'),
    (247, '如图所示，点O为四面体ABCD的外接球球心'),
]

PBDR = ('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="4" w:color="auto"/>'
        '<w:left w:val="single" w:sz="4" w:space="4" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="4" w:color="auto"/>'
        '<w:right w:val="single" w:sz="4" w:space="4" w:color="auto"/>'
        '</w:pBdr>')

doc = Document(SRC)
paras = doc.paragraphs
count = 0
for idx, prefix in BOXES:
    p = paras[idx]
    t = ''.join(x.text or '' for x in p._p.iter(WC + 't'))
    assert t.startswith(prefix), 'p%d 文本前缀不符: 期望%r 实际%r' % (idx, prefix, t[:20])
    ppr = p._p.get_or_add_pPr()
    assert ppr.find(qn('w:pBdr')) is None, 'p%d 已有pBdr' % idx
    # pBdr 须位于 pStyle/keepNext/…/numPr/suppressLineNumbers 之后、shd/spacing 之前
    PRE = ['pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr', 'widowControl',
           'numPr', 'suppressLineNumbers']
    anchor = None
    for child in ppr:
        name = child.tag.split('}')[-1]
        if name in PRE:
            anchor = child
        else:
            break
    if anchor is not None:
        anchor.addnext(parse_xml(PBDR))
    else:
        ppr.insert(0, parse_xml(PBDR))
    count += 1
assert count == len(BOXES) == 45, '框段数=%d≠45' % count
doc.save(DST)
print('N23定理框：%d 段加框（27条目/45框段）PASS -> %s' % (count, DST))
