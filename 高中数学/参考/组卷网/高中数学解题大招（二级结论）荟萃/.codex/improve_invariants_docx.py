from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_outline(paragraph, level: int | None) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:outlineLvl"))
    if old is not None:
        ppr.remove(old)
    if level is not None:
        node = OxmlElement("w:outlineLvl")
        node.set(qn("w:val"), str(level))
        ppr.append(node)


def set_keep_with_next(paragraph, enabled: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:keepNext"))
    if old is not None:
        ppr.remove(old)
    if enabled:
        ppr.append(OxmlElement("w:keepNext"))


def set_text(paragraph, text: str) -> None:
    paragraph.text = text


def append_text(parent, text: str, bold: bool = False) -> None:
    run = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        run.append(rpr)
    node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    parent.append(run)


def append_inline_math(parent, text: str) -> None:
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    node = OxmlElement("m:t")
    node.text = text
    run.append(node)
    math.append(run)
    parent.append(math)


def set_mixed(paragraph, segments: list[tuple[str, str]]) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    for kind, value in segments:
        if kind == "math":
            append_inline_math(p, value)
        else:
            append_text(p, value)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    target = Path(sys.argv[1]).resolve()
    backup_dir = target.parents[2] / ".codex" / "backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    backup = backup_dir / f"{target.stem}.修改前备份{target.suffix}"
    if not backup.exists():
        shutil.copy2(target, backup)

    doc = Document(target)
    original = list(doc.paragraphs)

    # 删去脱离题目的泛化知识讲解和重复的泛泛点评，保留纯题目驱动结构。
    for index in [3, 4, 5, 6, 7, 8, 9, 10, 20, 21, 22, 23, 44, 80, 171, 176]:
        delete_paragraph(original[index])

    replacements = {
        13: "【不变量】点 P 沿棱 DD₁ 运动时，点 P 到固定平面 ACC₁A₁ 的距离保持不变。于是三角形 PA₁C 的面积最小问题，可降为点 P 到平面内定直线 A₁C 的距离最小问题。",
        30: "【不变量】折起只改变图形的位置，不改变对应边长，因此点 P 到两个定点 B、D 的距离和保持不变；点 P 的轨迹可由这个定长和确定。",
        49: "【不变量】（1）点 D 沿棱 A₁B₁ 运动时，直线 DE 始终位于同一个固定垂面内；（2）三角形 DEF 在侧面 BB₁C₁C 上的投影始终是固定三角形 B₁GF，且线段 EF 的长度不变。",
        89: "【不变量】点 P 沿棱 AA₁ 运动时，底边 BC 固定，且直线 BC 始终垂直于线段 PE。因此三角形 PBC 的面积只随 PE 变化，可转化为点 P 到定点 E 的距离最小。",
        119: "【不变量】由中位线关系可以作出方向固定且平行于 EF 的直线 BG。点 P 运动时，原异面直线所成的角始终可转化为直线 BP 与固定直线 BG 所成的角。",
        143: "【不变量】线段 EF、PQ 的长度保持不变；平面 EFQ 始终是固定侧面 ACC₁A₁，点 M 沿与该平面平行的棱 BB₁ 运动，所以点 M 到平面 EFQ 的距离保持不变。",
        161: "【不变量】竖棱 CC′ 的长度固定，且它与动平面 PQC′ 所成的角固定，因此点 C 到动平面 PQC′ 的距离保持不变；同一四面体换底计算时，体积保持不变。",
        175: "【不变量】点 E 在线段 AB 上运动时，四棱锥的高、底面半边以及二面角 S-AB-C 的大小都不随点 E 改变，可把它作为固定比较角。",
        45: "【典例2】在直三棱柱 ABC-A₁B₁C₁ 中，侧面 AA₁B₁B 为正方形，棱 AB 与 BC 的长度均为 2。点 E、F 分别为 AC、CC₁ 的中点，点 D 在棱 A₁B₁ 上，直线 BF 垂直于棱 A₁B₁。",
        47: "（1）求证：直线 BF 垂直于直线 DE；",
        48: "（2）当 B₁D 为何值时，平面 BB₁C₁C 与平面 DFE 所成二面角的正弦值最小？",
        50: "（1）【法一】固定垂面：证明随动直线始终位于同一垂直平面",
        52: "取 BC 的中点 G，连接 EG、GB₁，如图①。",
        53: "因为 E、G 分别为 AC、BC 的中点，所以 EG 与 AB、A₁B₁ 平行，故 E、G、B₁、D 四点共面。",
        54: "四边形 BCC₁B₁ 为正方形，F、G 分别为 CC₁、BC 的中点，所以直线 BF 垂直于直线 GB₁。",
        55: "又因为直线 BF 垂直于棱 A₁B₁，直线 GB₁ 与棱 A₁B₁ 相交于点 B₁，且二者都在平面 EGB₁D 内，所以直线 BF 垂直于平面 EGB₁D。",
        56: "直线 DE 在平面 EGB₁D 内，因此直线 BF 垂直于直线 DE。",
        59: "因为三棱柱 ABC-A₁B₁C₁ 是直三棱柱，所以棱 BB₁ 垂直于底面 ABC，从而棱 BB₁ 垂直于棱 AB。",
        60: "棱 A₁B₁ 与棱 AB 平行，直线 BF 垂直于棱 A₁B₁，所以直线 BF 垂直于棱 AB。又因为直线 BB₁ 与 BF 相交于点 B，且都在侧面 BCC₁B₁ 内，所以棱 AB 垂直于侧面 BCC₁B₁，故 BA、BC、BB₁ 两两垂直。",
        73: "根据第（1）问的法二，设平面 DFE 的一个法向量为 ",
        75: "令 ",
        177: "【法一】固定比较角：在垂直截面内利用正切值比较",
        178: "点 E 在线段 AB 上运动时，四棱锥的高 SO、底面半边 OM 以及二面角 S-AB-C 的大小均为定值。以这个二面角为固定比较角，在垂直截面内分别进行比较。",
    }
    for index, text in replacements.items():
        set_text(original[index], text)

    set_mixed(
        original[61],
        [
            ("text", "以 B 为坐标原点，以 BA、BC、BB₁ 所在直线分别为 x 轴、y 轴、z 轴建立空间直角坐标系，如图②，则 "),
            ("math", "B(0,0,0)"), ("text", "，"),
            ("math", "A(2,0,0)"), ("text", "，"),
            ("math", "C(0,2,0)"), ("text", "，"),
            ("math", "B₁(0,0,2)"), ("text", "，"),
            ("math", "A₁(2,0,2)"), ("text", "，"),
            ("math", "C₁(0,2,2)"), ("text", "，"),
            ("math", "E(1,1,0)"), ("text", "，"),
            ("math", "F(0,2,1)"), ("text", "。由题意，设 "),
            ("math", "D(a,0,2)"), ("text", "，其中 "),
            ("math", "0≤a≤2"), ("text", "。"),
        ],
    )
    set_mixed(original[73], [("text", "根据第（1）问的法二，设平面 DFE 的一个法向量为 "), ("math", "m=(x,y,z)"), ("text", "。")])
    set_mixed(original[75], [("text", "令 "), ("math", "z=2−a"), ("text", "，则 "), ("math", "m=(3,1+a,2−a)"), ("text", "。")])
    set_mixed(original[167], [("text", "当 "), ("math", "ab=8"), ("text", " 时，"), ("math", "S△PQC=4"), ("text", "，棱锥 C′-PQC 的体积最小。")])

    # 统一“题目—答案—方法—分析”的大纲层级：答案折叠全部解法，各法可单独折叠。
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("【典例") or re.match(r"^[1-9]\d*．", text):
            set_outline(paragraph, 0)
            set_keep_with_next(paragraph)
        elif text.startswith("【答案】") or text.startswith("答案"):
            set_outline(paragraph, 1)
            set_keep_with_next(paragraph)
        elif "【法一】" in text or "【法二】" in text:
            set_outline(paragraph, 2)
            set_keep_with_next(paragraph)
        elif text.startswith("【分析】") or text.startswith("【详解】"):
            set_outline(paragraph, 3)
            set_keep_with_next(paragraph)

    doc.save(target)
    print(f"saved={target}")
    print(f"backup={backup}")
    print(f"paragraphs={len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
