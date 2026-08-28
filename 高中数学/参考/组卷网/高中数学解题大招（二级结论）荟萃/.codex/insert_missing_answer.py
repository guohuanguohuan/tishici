from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def main() -> None:
    target = Path(sys.argv[1]).resolve()
    doc = Document(target)
    if any("【答案】（1）证明见解析；（2）B₁D 为二分之一" in p.text for p in doc.paragraphs):
        return

    anchor = next(p for p in doc.paragraphs if p.text.strip().startswith("（2）当 B₁D 为何值时"))
    paragraph = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "1")
    ppr.append(outline)
    ppr.append(OxmlElement("w:keepNext"))
    paragraph.append(ppr)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:b"))
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "【答案】（1）证明见解析；（2）B₁D 为二分之一。"
    run.append(text)
    paragraph.append(run)
    anchor._p.addnext(paragraph)

    doc.save(target)


if __name__ == "__main__":
    main()
