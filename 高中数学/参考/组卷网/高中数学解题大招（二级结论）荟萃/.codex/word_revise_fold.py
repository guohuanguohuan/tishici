from __future__ import annotations

import argparse
import copy
import math
import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def local_name(node) -> str:
    return etree.QName(node).localname


def node_text(node) -> str:
    return "".join((n.text or "") for n in node.iter() if local_name(n) == "t")


def body_children(doc: Document):
    return list(doc._element.body)


def body_paragraphs(doc: Document):
    return [n for n in body_children(doc) if local_name(n) == "p"]


def first_para(doc: Document, predicate, start_after=None):
    passed = start_after is None
    for p in body_paragraphs(doc):
        if not passed:
            if p is start_after:
                passed = True
            continue
        if predicate(node_text(p).strip()):
            return p
    raise RuntimeError("paragraph not found")


def p_starts(doc: Document, prefix: str):
    return first_para(doc, lambda t: t.startswith(prefix))


def paragraphs_between(doc: Document, start, end_exclusive):
    children = body_children(doc)
    i = children.index(start)
    j = children.index(end_exclusive) if end_exclusive is not None else len(children) - 1
    return [n for n in children[i:j] if local_name(n) == "p"]


def next_problem_start(doc: Document, start):
    children = body_children(doc)
    i = children.index(start)
    for n in children[i + 1 :]:
        if local_name(n) != "p":
            continue
        t = node_text(n).strip()
        if t.startswith("【典例") or re.match(r"^\d+[．.]", t):
            return n
    return None


def make_paragraph(text: str, bold=False, outline=None):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    if outline is not None:
        el = OxmlElement("w:outlineLvl")
        el.set(qn("w:val"), str(outline))
        pPr.append(el)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "1")
        bCs = OxmlElement("w:bCs")
        bCs.set(qn("w:val"), "1")
        rPr.extend([b, bCs])
        r.append(rPr)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_before(ref, new_node):
    ref.getparent().insert(ref.getparent().index(ref), new_node)


def insert_after(ref, new_node):
    ref.getparent().insert(ref.getparent().index(ref) + 1, new_node)


def replace_text_once(p, old: str, new: str) -> bool:
    for n in p.iter():
        if local_name(n) == "t" and n.text and old in n.text:
            n.text = n.text.replace(old, new, 1)
            return True
    return False


def replace_text_all(p, replacements):
    for n in p.iter():
        if local_name(n) == "t" and n.text:
            for old, new in replacements:
                n.text = n.text.replace(old, new)


def set_outline(p, level: int):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    old = pPr.find(qn("w:outlineLvl"))
    if old is not None:
        pPr.remove(old)
    el = OxmlElement("w:outlineLvl")
    el.set(qn("w:val"), str(level))
    pPr.append(el)


def bold_paragraph(p):
    for r in p.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)
        for tag in ("w:b", "w:bCs"):
            if rPr.find(qn(tag)) is None:
                e = OxmlElement(tag)
                e.set(qn("w:val"), "1")
                rPr.append(e)


def remove_blank_neighbors(node):
    parent = node.getparent()
    idx = parent.index(node)
    for direction in (-1, 1):
        j = idx + direction
        while 0 <= j < len(parent):
            n = parent[j]
            if local_name(n) != "p" or node_text(n).strip() or any(local_name(x) in ("drawing", "pict") for x in n.iter()):
                break
            parent.remove(n)
            if direction < 0:
                idx -= 1
                j -= 1


def unwrap_display_math(doc: Document):
    for p in list(doc._element.iter(qn("w:p"))):
        for omp in list(p.iter()):
            if local_name(omp) != "oMathPara":
                continue
            parent = omp.getparent()
            idx = parent.index(omp)
            # A text run keeps Word from promoting a formula-only paragraph
            # back to display math on save; ∴ also fits the continuation line.
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = "∴"
            r.append(t)
            parent.insert(idx, r)
            idx += 1
            children = [ch for ch in list(omp) if local_name(ch) == "oMath"]
            for ch in children:
                omp.remove(ch)
                parent.insert(idx, ch)
                idx += 1
            parent.remove(omp)


def paragraph_weight(p) -> float:
    text = node_text(p).strip()
    weight = max(0.7, math.ceil(max(1, len(text)) / 25))
    for n in p.iter():
        if local_name(n) == "extent" and etree.QName(n).namespace == WP_NS:
            try:
                height_pt = int(n.get("cy")) / 12700
                weight += max(4, height_pt / 15)
            except Exception:
                weight += 8
    return weight


def remove_default_cell_para(tc):
    for p in list(tc):
        if local_name(p) == "p" and not node_text(p).strip() and not any(local_name(x) in ("drawing", "pict") for x in p.iter()):
            tc.remove(p)


def set_cell_width(tc, width_twips):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "top")
    tcPr.append(vAlign)


def configure_table(tbl, text_width_twips):
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for tag in ("w:tblW", "w:tblLayout", "w:tblBorders", "w:tblCellMar"):
        old = tblPr.find(qn(tag))
        if old is not None:
            tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), "72")
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for c in list(grid):
        grid.remove(c)
    half = text_width_twips // 2
    for _ in range(2):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(half))
        grid.append(col)
    row = tbl.find(qn("w:tr"))
    cells = row.findall(qn("w:tc"))
    for tc in cells:
        set_cell_width(tc, half)


def make_two_col_table(doc: Document, analysis_paras, method_name: str, step_specs, text_width_twips):
    if not analysis_paras:
        raise RuntimeError("empty analysis")

    # Add semantic step titles immediately before matched complete paragraphs.
    for match, title in reversed(step_specs):
        target = next((p for p in analysis_paras if match in node_text(p)), None)
        if target is not None:
            step_p = make_paragraph(title, bold=True)
            target.addprevious(step_p)
            analysis_paras.insert(analysis_paras.index(target), step_p)

    table = doc.add_table(rows=1, cols=2)
    tbl = table._tbl
    configure_table(tbl, text_width_twips)
    insert_before(analysis_paras[0], tbl)
    row = tbl.find(qn("w:tr"))
    left, right = row.findall(qn("w:tc"))
    remove_default_cell_para(left)
    remove_default_cell_para(right)

    heading = make_paragraph("【详细解析】", bold=True, outline=2)
    method = make_paragraph(f"【法一：{method_name}】", bold=True, outline=3)
    left.append(heading)
    left.append(method)

    # Balance by rendered-height proxy; only split at whole paragraph boundaries.
    weights = [paragraph_weight(p) for p in analysis_paras]
    prefix = []
    s = 0.0
    for w in weights:
        s += w
        prefix.append(s)
    header_weight = 2.0
    best_i = 1 if len(analysis_paras) > 1 else len(analysis_paras)
    best_diff = float("inf")
    for i in range(1, len(analysis_paras)):
        lw = header_weight + prefix[i - 1]
        rw = prefix[-1] - prefix[i - 1]
        diff = abs(lw - rw)
        if diff < best_diff:
            best_diff, best_i = diff, i
    for i, p in enumerate(analysis_paras):
        (left if i < best_i else right).append(p)
    if not any(local_name(n) == "p" for n in right):
        right.append(make_paragraph(""))
    return tbl


def set_paragraph_layout(p):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "left")
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    snap = pPr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        pPr.append(snap)
    snap.set(qn("w:val"), "0")


def set_all_layout(doc: Document, only_after_text=None):
    active = only_after_text is None
    for p in doc._element.body.iter(qn("w:p")):
        if not active and only_after_text in node_text(p):
            active = True
        if active:
            set_paragraph_layout(p)


def text_width_twips(doc: Document):
    sec = doc.sections[0]
    return int((sec.page_width - sec.left_margin - sec.right_margin) / 635)  # EMU -> twips


SOURCE_SPECS = [
    {
        "start": "【典例1】",
        "tag": "翻折周长问题",
        "idea": "【答案思路】侧面展开，⟦A,E,F,A_1⟧ 共线时 ⟦C_{△AEF,min}=AA_1⟧；再由余弦定理求 ⟦AA_1⟧。",
        "method": "翻折平面化",
        "steps": [("【解析】", "【步骤1 展开取直】")],
        "example": True,
    },
    {
        "start": "1．在正三棱锥",
        "tag": "翻折周长问题",
        "idea": "【答案思路】侧面展开，⟦A,D,E,A'⟧ 共线时 ⟦C_{△ADE,min}=AA'⟧；由 ⟦sin(α/2)=1/4⟧ 求 ⟦sin(3α/2)⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开侧面】"), ("则当", "【步骤2 共线取小】"), ("设∠APB", "【步骤3 计算距离】")],
    },
    {
        "start": "【典例2】",
        "tag": "翻折距离问题",
        "idea": "【答案思路】连续展开 ⟦△PAD,△PCD,△PBC⟧，则 ⟦min(AE+EF+BF)=AB'⟧；再由余弦定理求 ⟦AB'⟧。",
        "method": "翻折平面化",
        "steps": [("【解析】", "【步骤1 证明垂直】"), ("将平面", "【步骤2 连续展开】"), ("∵cos", "【步骤3 余弦计算】")],
        "example": True,
    },
    {
        "start": "2．如图，在四棱锥",
        "tag": "翻折距离问题",
        "idea": "【答案思路】由 ⟦BC⊥平面PAB⟧，沿 ⟦BC⟧ 翻折 ⟦△PBC⟧；⟦P',E,D⟧ 共线时 ⟦PE+DE⟧ 最小。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 证折痕垂直】"), ("将△PBC", "【步骤2 翻折取小】")],
    },
    {
        "start": "3．已知正三棱锥",
        "tag": "翻折周长问题",
        "idea": "【答案思路】侧面展开，⟦A,D,E,A_1⟧ 共线时 ⟦C_{△ADE,min}=AA_1⟧；由 ⟦∠APA_1=90°⟧ 求 ⟦AA_1⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开侧面】"), ("则将求", "【步骤2 共线取小】"), ("结合题意", "【步骤3 计算距离】")],
    },
    {
        "start": "4．如图，O为正四棱锥",
        "tag": "翻折距离问题",
        "idea": "【答案思路】⟦△POC≌△POD⟧，将 ⟦△POD⟧ 绕 ⟦PO⟧ 翻至 ⟦△POC⟧，则原折线最小值转为 ⟦d(A,PC)⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 找全等面】"), ("将△POD", "【步骤2 翻折转距离】"), ("又因为", "【步骤3 计算点线距】")],
    },
    {
        "start": "5．在四棱锥",
        "tag": "翻折周长问题",
        "idea": "【答案思路】沿 ⟦PA⟧ 展开四个侧面，⟦A,E,F,G,A'⟧ 共线时 ⟦C_{AEFG,min}=AA'⟧，再由余弦定理求值。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开侧面】"), ("AE+EF", "【步骤2 共线取小】"), ("因为PA", "【步骤3 余弦计算】")],
    },
    {
        "start": "6．正三棱锥",
        "tag": "翻折周长问题",
        "idea": "【答案思路】剪开侧面得 ⟦SABCA'⟧，⟦A,E,F,A'⟧ 共线时 ⟦C_{△AEF,min}=AA'⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开侧面】"), ("连接AA'", "【步骤2 共线取小】"), ("正三棱锥", "【步骤3 计算距离】")],
    },
    {
        "start": "7．在意大利",
        "tag": "翻折圆锥体积问题",
        "idea": "【答案思路】展开圆锥侧面，⟦A'C⟧ 为最短灯带；由余弦定理求扇形圆心角，再由 ⟦l=Rθ=2πr⟧ 求 ⟦r,h,V⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开圆锥】"), ("△A'SC", "【步骤2 余弦定角】"), ("所以∠", "【步骤3 回代体积】")],
    },
    {
        "start": "8．如图，圆锥",
        "tag": "翻折圆锥距离问题",
        "idea": "【答案思路】展开半个圆锥，⟦BD_1⟧ 为最短路；由 ⟦AB^2+AD_1^2=BD_1^2⟧ 求扇形角，再由弧长求 ⟦r⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开半圆锥】"), ("连接BD", "【步骤2 确定最短线】"), ("设展开图", "【步骤3 弧长求半径】")],
    },
    {
        "start": "9．如图，三棱柱",
        "tag": "翻折棱柱距离问题",
        "idea": "【答案思路】沿 ⟦BC_1⟧ 展开相邻面，⟦C,P,A_1⟧ 共线时 ⟦min(CP+PA_1)=CA_1⟧；再用余弦定理。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开相邻面】"), ("连接A1C", "【步骤2 共线取小】"), ("因为∠ACB", "【步骤3 余弦计算】")],
    },
    {
        "start": "10．如图，在长方体",
        "tag": "翻折长方体距离问题",
        "idea": "【答案思路】展开相邻矩形，⟦A,M,D_1⟧ 共线时折线最短；再将原空间条件 ⟦MD_1⊥MA⟧ 转为勾股关系求 ⟦CM⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 展开矩形】"), ("当A,M", "【步骤2 共线建比例】"), ("原图形中", "【步骤3 垂直列方程】")],
    },
    {
        "start": "11．如图，直三棱柱",
        "tag": "翻折棱柱距离问题",
        "idea": "【答案思路】比较两种展开图，仅保留线段穿过指定棱的方案；在可行展开图中，最短路 ⟦=EF⟧。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 比较展开】"), ("若将底面ABC沿BC", "【步骤2 筛选路径】"), ("过E作", "【步骤3 计算最短路】")],
    },
    {
        "start": "12．在《九章算术》",
        "tag": "翻折外接球问题",
        "idea": "【答案思路】⟦PF⟧ 固定，故只需 ⟦min(PE+EF)⟧；展开侧面并令 ⟦P,E,F⟧ 共线，再由外接圆半径求外接球半径。",
        "method": "翻折平面化",
        "steps": [("【详解】", "【步骤1 固定周长项】"), ("如图，将四棱锥", "【步骤2 翻折取小】"), ("易得", "【步骤3 计算外接球】")],
    },
]


TARGET_SPECS = [
    {
        "start": "15．已知正方体",
        "tag": "截面形状问题",
        "idea": "【答案思路】以截面由四边形变为五边形的临界位置为界；由 ⟦MN∥AD_1⟧ 得临界值，再判 ⟦BM⟧ 的范围。",
        "method": "补全截面",
        "steps": [("【详解】", "【步骤1 还原棱长】"), ("如图所示", "【步骤2 找临界面】"), ("所以，当", "【步骤3 判定范围】")],
    },
    {
        "start": "16．如图，四棱柱",
        "tag": "截面综合问题",
        "idea": "【答案思路】①判体积不变量；②由平行关系补全截面；③算截面周长；④由距离比反推交点位置，逐项判定。",
        "method": "补全截面",
        "steps": [("【详解】", "【步骤1 判断体积】"), ("对于②", "【步骤2 补全截面】"), ("对于③", "【步骤3 计算周长】"), ("对于④", "【步骤4 检验距离】")],
    },
    {
        "start": "17．已知正方体",
        "tag": "截面形状问题",
        "idea": "【答案思路】先取 ⟦MN∥AD_1⟧ 的临界值，再按 ⟦CM⟧ 与临界值的大小分类作截面，确定五边形对应区间。",
        "method": "补全截面",
        "steps": [("【详解】", "【步骤1 找临界值】"), ("当CM", "【步骤2 分类作图】"), ("综上所述", "【步骤3 确定范围】")],
    },
]


def problem_analysis_paras(doc: Document, start, example=False):
    end = next_problem_start(doc, start)
    unit = paragraphs_between(doc, start, end)
    if example:
        a = next((p for p in unit if node_text(p).strip().startswith("【解析】")), None)
        stop = next((p for p in unit if node_text(p).strip().startswith("【题后反思】")), None)
    else:
        a = next((p for p in unit if node_text(p).strip().startswith("【分析】")), None)
        stop = next((p for p in unit if node_text(p).strip().startswith("【点睛】")), None)
    if a is None:
        raise RuntimeError(f"analysis start missing: {node_text(start)[:50]}")
    children = body_children(doc)
    i = children.index(a)
    if stop is not None:
        j = children.index(stop)
    elif end is not None:
        j = children.index(end)
    else:
        j = len(children) - 1
    paras = [n for n in children[i:j] if local_name(n) == "p"]
    # Exclude trailing pure blanks.
    while paras and not node_text(paras[-1]).strip() and not any(local_name(x) in ("drawing", "pict") for x in paras[-1].iter()):
        paras.pop()
    return paras


def find_standard_answer(unit, example=False):
    if example:
        for p in unit:
            t = node_text(p).strip()
            if t.startswith("故答案为") or t.startswith("故选"):
                return p
    else:
        for p in unit:
            if node_text(p).strip().startswith("【答案】"):
                return p
    raise RuntimeError("standard answer missing")


def normalize_answer(answer_p, example=False):
    if example:
        if not replace_text_once(answer_p, "故答案为：", "【标准答案】"):
            if not replace_text_once(answer_p, "故答案为:", "【标准答案】"):
                replace_text_once(answer_p, "故选：", "【标准答案】")
    else:
        replace_text_once(answer_p, "【答案】", "【标准答案】")
    set_outline(answer_p, 2)
    bold_paragraph(answer_p)


def format_problem(doc: Document, spec, text_width):
    start = p_starts(doc, spec["start"])
    end = next_problem_start(doc, start)
    unit = paragraphs_between(doc, start, end)
    example = spec.get("example", False)

    label = make_paragraph(f"【题型：{spec['tag']}】", bold=True, outline=1)
    insert_before(start, label)

    answer = find_standard_answer(unit, example=example)
    normalize_answer(answer, example=example)

    if example:
        idea_existing = next((p for p in unit if node_text(p).strip().startswith("【大招指引】")), None)
        if idea_existing is not None:
            replace_text_once(idea_existing, "【大招指引】", "【答案思路】")
            set_outline(idea_existing, 2)
            bold_paragraph(idea_existing)
            # Put the standard answer before the answer idea.
            idea_existing.addprevious(answer)
        else:
            idea_p = make_paragraph(spec["idea"], bold=True, outline=2)
            insert_after(answer, idea_p)
    else:
        idea_p = make_paragraph(spec["idea"], bold=True, outline=2)
        insert_after(answer, idea_p)

    analysis = problem_analysis_paras(doc, start, example=example)
    for p in analysis:
        replace_text_all(p, [("又因为", "又∵"), ("因为", "∵"), ("因此", "∴"), ("所以", "∴")])
    make_two_col_table(doc, analysis, spec["method"], spec["steps"], text_width)


def replace_intro(doc: Document):
    p1 = p_starts(doc, "处理不同平面内")
    p2 = p_starts(doc, "立体几何是高考")
    new_paras = [
        make_paragraph("【符号约定】设 ⟦α∩β=l⟧；绕 ⟦l⟧ 将 ⟦β⟧ 翻至 ⟦α⟧，记 ⟦T_l:β→α，P↦P'⟧。", bold=False),
        make_paragraph("【翻折不变量】⟦∀M∈l，T_l(M)=M⟧；若 ⟦P↦P'，Q↦Q'⟧，则 ⟦MP=MP'，PQ=P'Q'，∠PMQ=∠P'MQ'⟧。", bold=False),
        make_paragraph("【平面化】若 ⟦A∈α，B∈β，M∈l，B'=T_l(B)⟧，则 ⟦AM+MB=AM+MB'≥AB'⟧；等号成立 ⟦⇔⟧ ⟦A,M,B'⟧ 共线。", bold=False),
        make_paragraph("【最值模型】⟦空间折线→平面折线→平面线段⟧；⟦min_(M∈l)(AM+MB)=AB'⟧。周长最值同理：连续展开后，⟦C_min=⟧ 两个对应端点间的线段长。", bold=False),
    ]
    parent = p1.getparent()
    idx = parent.index(p1)
    parent.remove(p1)
    if p2.getparent() is parent:
        parent.remove(p2)
    for k, p in enumerate(new_paras):
        parent.insert(idx + k, p)


def add_terms_summary(doc: Document, terms, marker=None):
    body = doc._element.body
    sect = body.find(qn("w:sectPr"))
    p = make_paragraph("【题型术语】" + "、".join(terms), bold=True, outline=1)
    if sect is not None:
        body.insert(body.index(sect), p)
    else:
        body.append(p)


def format_source(path: Path):
    doc = Document(path)
    unwrap_display_math(doc)
    replace_intro(doc)
    width = text_width_twips(doc)
    for spec in SOURCE_SPECS:
        format_problem(doc, spec, width)
    terms = []
    for s in SOURCE_SPECS:
        if s["tag"] not in terms:
            terms.append(s["tag"])
    add_terms_summary(doc, terms)
    set_all_layout(doc)
    doc.save(path)


def format_target(path: Path):
    doc = Document(path)
    unwrap_display_math(doc)
    width = text_width_twips(doc)
    for spec in TARGET_SPECS:
        format_problem(doc, spec, width)
    add_terms_summary(doc, ["截面形状问题", "截面综合问题"])
    set_all_layout(doc, only_after_text="【由翻折专题转入】")
    doc.save(path)


def find_range(doc, text: str, start=0):
    r = doc.Range(Start=start, End=doc.Content.End)
    f = r.Find
    f.ClearFormatting()
    f.Text = text
    f.Forward = True
    f.Wrap = 0
    if not f.Execute():
        raise RuntimeError(f"Word text not found: {text}")
    return r


def prepare_source_with_word(source: Path, temp_dir: Path):
    import win32com.client as win32

    word = None
    doc = None
    temp_dir.mkdir(parents=True, exist_ok=True)
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(source), ReadOnly=False, AddToRecentFiles=False)
        ranges = [
            ("4．已知正方体", "5．如图", "moved_q4.docx"),
            ("7．如图，四棱柱", "8．已知正方体", "moved_q7.docx"),
            ("8．已知正方体", "9．正三棱锥", "moved_q8.docx"),
        ]
        located = []
        for start_text, end_text, name in ranges:
            s = find_range(doc, start_text).Start
            e = find_range(doc, end_text, start=s + 1).Start
            located.append((s, e, temp_dir / name))
        for s, e, out in located:
            td = word.Documents.Add()
            td.Content.FormattedText = doc.Range(s, e).FormattedText
            td.SaveAs2(str(out), FileFormat=16)
            td.Close(False)
        for s, e, _ in sorted(located, reverse=True):
            doc.Range(s, e).Delete()

        # Renumber remaining exercises while preserving all formula/image objects.
        mappings = [
            ("5．如图", "4．"),
            ("6．在四棱锥", "5．"),
            ("9．正三棱锥", "6．"),
            ("10．在意大利", "7．"),
            ("11．如图，圆锥", "8．"),
            ("12．如图，三棱柱", "9．"),
            ("13．如图，在长方体", "10．"),
            ("14．如图，直三棱柱", "11．"),
            ("15．在《九章算术》", "12．"),
        ]
        for old, new in mappings:
            r = find_range(doc, old)
            dot = old.index("．") + 1
            doc.Range(r.Start, r.Start + dot).Text = new
        doc.Save()
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def prepare_target_with_word(target: Path, temp_dir: Path):
    import win32com.client as win32

    word = None
    doc = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(target), ReadOnly=False, AddToRecentFiles=False)
        end = doc.Content.End - 1
        r = doc.Range(end, end)
        r.InsertAfter("\r【由翻折专题转入】\r")
        r = doc.Range(doc.Content.End - 1, doc.Content.End - 1)
        # InsertFile at a collapsed end range stacks subsequent insertions at
        # the same anchor, so reverse here to obtain q4 -> q7 -> q8 visually.
        for name in ("moved_q8.docx", "moved_q7.docx", "moved_q4.docx"):
            r.InsertFile(str(temp_dir / name))
            r.Collapse(0)
            r.InsertAfter("\r")
            r.Collapse(0)
        # Renumber only the appended blocks, found after the transfer heading.
        marker = find_range(doc, "【由翻折专题转入】").End
        for old, new in (("4．已知正方体", "15．"), ("7．如图，四棱柱", "16．"), ("8．已知正方体", "17．")):
            rr = find_range(doc, old, start=marker)
            dot = old.index("．") + 1
            doc.Range(rr.Start, rr.Start + dot).Text = new
        doc.Save()
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def convert_markers_to_word_math(doc):
    # Converts ⟦...⟧ markers to real inline Microsoft OMath objects.
    while True:
        start_r = doc.Range(0, doc.Content.End)
        start_r.Find.Text = "⟦"
        start_r.Find.Forward = True
        start_r.Find.Wrap = 0
        if not start_r.Find.Execute():
            break
        end_r = doc.Range(start_r.End, doc.Content.End)
        end_r.Find.Text = "⟧"
        end_r.Find.Forward = True
        end_r.Find.Wrap = 0
        if not end_r.Find.Execute():
            raise RuntimeError("unclosed math marker")
        start = start_r.Start
        length = end_r.Start - start_r.End
        end_r.Delete()
        start_r.Delete()
        mr = doc.Range(start, start + length)
        doc.OMaths.Add(mr)
        if mr.OMaths.Count:
            mr.OMaths(1).BuildUp()


def word_finalize(path: Path, pdf_path: Path | None = None):
    import win32com.client as win32

    word = None
    doc = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        word.Options.ConfirmConversions = False
        doc = word.Documents.Open(str(path), ReadOnly=False, AddToRecentFiles=False)
        convert_markers_to_word_math(doc)
        # Enforce exact table width after Word repagination.
        for table in doc.Tables:
            table.AllowAutoFit = False
            table.Rows.AllowBreakAcrossPages = False
            table.PreferredWidthType = 2  # wdPreferredWidthPercent
            table.PreferredWidth = 100
            if table.Columns.Count == 2:
                table.Columns(1).PreferredWidthType = 2
                table.Columns(1).PreferredWidth = 50
                table.Columns(2).PreferredWidthType = 2
                table.Columns(2).PreferredWidth = 50
            table.Borders.Enable = True
        doc.Repaginate()
        doc.Save()
        if pdf_path is not None:
            doc.ExportAsFixedFormat(str(pdf_path), 17)
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def validate_docx(path: Path, expected_tables: int, expected_labels: int, marker=None):
    doc = Document(path)
    root = doc._element
    text = node_text(root)
    errors = []
    if "⟦" in text or "⟧" in text:
        errors.append("unconverted math markers")
    display = sum(1 for n in root.iter() if local_name(n) == "oMathPara")
    if display:
        errors.append(f"display equations={display}")
    tables = len(doc.tables)
    if tables < expected_tables:
        errors.append(f"tables={tables} expected>={expected_tables}")
    labels = text.count("【题型：")
    if labels < expected_labels:
        errors.append(f"labels={labels} expected>={expected_labels}")
    ideas = text.count("【答案思路】")
    if ideas < expected_labels:
        errors.append(f"ideas={ideas} expected>={expected_labels}")
    methods = text.count("【法一：")
    if methods < expected_labels:
        errors.append(f"methods={methods} expected>={expected_labels}")
    for table in doc.tables[-expected_tables:]:
        if len(table.columns) != 2:
            errors.append("non-two-column detail table")
    return {
        "file": str(path),
        "paragraphs": len(doc.paragraphs),
        "tables": tables,
        "labels": labels,
        "ideas": ideas,
        "methods": methods,
        "inline_math": sum(1 for n in root.iter() if local_name(n) == "oMath"),
        "display_math": display,
        "errors": errors,
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("action", choices=["prepare-source", "format-source", "prepare-target", "format-target", "finalize", "validate"])
    ap.add_argument("--source", type=Path)
    ap.add_argument("--target", type=Path)
    ap.add_argument("--temp-dir", type=Path)
    ap.add_argument("--pdf", type=Path)
    args = ap.parse_args()
    if args.action == "prepare-source":
        prepare_source_with_word(args.source.resolve(), args.temp_dir.resolve())
    elif args.action == "format-source":
        format_source(args.source.resolve())
    elif args.action == "prepare-target":
        prepare_target_with_word(args.target.resolve(), args.temp_dir.resolve())
    elif args.action == "format-target":
        format_target(args.target.resolve())
    elif args.action == "finalize":
        word_finalize(args.source.resolve(), args.pdf.resolve() if args.pdf else None)
    elif args.action == "validate":
        if args.source:
            print(validate_docx(args.source.resolve(), 14, 14))
        if args.target:
            print(validate_docx(args.target.resolve(), 3, 3))


if __name__ == "__main__":
    main()
