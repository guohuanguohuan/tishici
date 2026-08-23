from pathlib import Path
import sys

from docx import Document


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: edit_keep_proof1.py INPUT.docx OUTPUT.docx")

    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document(source)

    start_marker = "【证明二：垂直截面几何法】"
    end_marker = "二、空间正弦定理（三面角正弦定理）"

    start = next((i for i, p in enumerate(doc.paragraphs) if p.text.startswith(start_marker)), None)
    end = next((i for i, p in enumerate(doc.paragraphs) if p.text.startswith(end_marker)), None)
    if start is None or end is None or start >= end:
        raise RuntimeError(f"unexpected proof range: start={start}, end={end}")

    # Remove only proof two, including its equations and illustration. Keep the
    # existing blank separator immediately before it so section two still has
    # natural breathing room after proof one.
    elements = [p._p for p in doc.paragraphs[start:end]]
    for element in elements:
        element.getparent().remove(element)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    check = Document(output)
    body_text = "\n".join(p.text for p in check.paragraphs)
    if start_marker in body_text:
        raise RuntimeError("proof two marker remains in output")
    if "【证明一：向量投影法】" not in body_text or end_marker not in body_text:
        raise RuntimeError("required content is missing after edit")

    print(f"removed_paragraphs={len(elements)}")
    print(f"output={output}")


if __name__ == "__main__":
    main()
