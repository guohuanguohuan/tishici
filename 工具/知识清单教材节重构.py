# -*- coding: utf-8 -*-
"""按目标教材目录重构数学/物理知识清单的节标题。

仅调整结构标题与挂靠位置，不改题目、答案、公式、表格或图片。
用法: python 工具/知识清单教材节重构.py <math|physics|physics-final> <输入.docx> <输出.docx>
  math=删「一、二、三」源结构标题并按选必2第3章目录插节标题；physics=插必修3第12章四节标题并删旧条目标题；physics-final=仅把12.3标题移回正确挂靠并把文内/core title改为输出文件名。
"""
from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def paragraph_text(p_el):
    return "".join(t.text or "" for t in p_el.findall(".//" + W + "t"))


def body_paragraphs(doc):
    return [el for el in doc.element.body if el.tag == W + "p"]


def find_para(doc, prefix):
    for p_el in body_paragraphs(doc):
        if paragraph_text(p_el).strip().startswith(prefix):
            return p_el
    raise ValueError(f"未找到段落：{prefix}")


def find_first_image_para(doc):
    """定位正文中的首个图片段落，供知识架构图挂靠教材节。"""
    for p_el in body_paragraphs(doc):
        if p_el.findall(".//" + W + "drawing") or p_el.findall(".//" + W + "pict"):
            return p_el
    raise ValueError("未找到图片段落")


def remove_para(p_el):
    parent = p_el.getparent()
    if parent is not None:
        parent.remove(p_el)


def make_heading(text, level):
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    p.append(ppr)
    jc = OxmlElement("w:jc")
    jc.set(W + "val", "left")
    ppr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(W + "before", "0")
    spacing.set(W + "after", "0")
    spacing.set(W + "line", "300")
    spacing.set(W + "lineRule", "auto")
    ppr.append(spacing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(W + "val", str(level))
    ppr.append(outline)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    r.append(rpr)
    fonts = OxmlElement("w:rFonts")
    fonts.set(W + "ascii", "Times New Roman")
    fonts.set(W + "hAnsi", "Times New Roman")
    fonts.set(W + "eastAsia", "宋体")
    rpr.append(fonts)
    for name in ("b", "bCs"):
        el = OxmlElement("w:" + name)
        el.set(W + "val", "1")
        rpr.append(el)
    for name in ("sz", "szCs"):
        el = OxmlElement("w:" + name)
        el.set(W + "val", "21")
        rpr.append(el)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_headings_before(anchor, headings):
    parent = anchor.getparent()
    idx = parent.index(anchor)
    for text, level in headings:
        parent.insert(idx, make_heading(text, level))
        idx += 1


def restructure_math(src, dst):
    doc = Document(src)
    for prefix in ("一、单元学习目标", "二、单元知识架构", "三、单元知识梳理"):
        remove_para(find_para(doc, prefix))
    # 知识架构图属于本章正文，按最新版图片规则挂到首个相关教材节下，
    # 不再作为章首装饰；图片本身保持原样。
    insert_headings_before(find_first_image_para(doc), [("3.1 排列与组合", 0)])
    insert_headings_before(find_para(doc, "1．教材梳理填空"), [("3.1.1 基本计数原理", 1)])
    insert_headings_before(find_para(doc, "5．排列："), [("3.1.2 排列与排列数", 1)])
    insert_headings_before(find_para(doc, "11．组合的定义"), [("3.1.3 组合与组合数", 1)])
    insert_headings_before(
        find_para(doc, "14．教材梳理填空"),
        [("3.2 数学探究活动：生日悖论的解释与模拟", 0),
         ("3.3 二项式定理与杨辉三角", 0)],
    )
    doc.save(dst)


def restructure_physics(src, dst):
    doc = Document(src)
    old_headings = [
        "1．电功和电功率", "2．焦耳定律", "3．电路中的能量转化", "4．电动势",
        "5．闭合电路欧姆定律", "6．路端电压与负载的关系", "7．电源的U-I图像",
        "8．能量守恒定律", "9．能量转移或转化的方向性",
        "10．能源的分类与应用　能源与社会发展",
    ]
    insert_headings_before(find_para(doc, old_headings[0]), [("12.1 电路中的能量转化", 0)])
    insert_headings_before(find_para(doc, old_headings[3]), [("12.2 闭合电路的欧姆定律", 0)])
    # “电源的 U-I 图像”是教材 12.3 测量电动势和内阻的数据处理方法，
    # 不能继续挂在 12.2，更不能把 12.3 留成空标题。
    insert_headings_before(
        find_para(doc, old_headings[6]),
        [("12.3 实验：电池电动势和内阻的测量", 0)],
    )
    insert_headings_before(find_para(doc, old_headings[7]), [("12.4 能源与可持续发展", 0)])
    for prefix in old_headings:
        remove_para(find_para(doc, prefix))
    doc.save(dst)


def replace_paragraph_text(p_el, text):
    runs = p_el.findall("./" + W + "r")
    if not runs:
        r = OxmlElement("w:r")
        p_el.append(r)
        runs = [r]
    first = runs[0]
    for child in list(first):
        if child.tag != W + "rPr":
            first.remove(child)
    t = OxmlElement("w:t")
    t.text = text
    first.append(t)
    for run in runs[1:]:
        p_el.remove(run)


def fix_physics_final(src, dst):
    """修正已填答成品的12.3挂靠及总控3.5完整命名。"""
    doc = Document(src)
    heading = find_para(doc, "12.3 实验：电池电动势和内阻的测量")
    anchor = find_para(doc, "（1）公式：U =")
    heading.getparent().remove(heading)
    anchor.addprevious(heading)
    expected_title = Path(dst).stem
    first = next(p for p in body_paragraphs(doc) if paragraph_text(p).strip())
    replace_paragraph_text(first, expected_title)
    doc.core_properties.title = expected_title
    doc.save(dst)


def main():
    if len(sys.argv) != 4 or sys.argv[1] not in ("math", "physics", "physics-final"):
        raise SystemExit("用法：python 知识清单教材节重构.py math|physics|physics-final 输入.docx 输出.docx")
    mode, src, dst = sys.argv[1], sys.argv[2], sys.argv[3]
    Path(dst).parent.mkdir(parents=True, exist_ok=True)
    if mode == "math":
        restructure_math(src, dst)
    elif mode == "physics":
        restructure_physics(src, dst)
    else:
        fix_physics_final(src, dst)
    print(f"OK {mode}: {src} -> {dst}")


if __name__ == "__main__":
    main()
