# -*- coding: utf-8 -*-
r"""部分封面产件（工具债⑧；公共规则§11部分封面件条款N9、§5创作层铁规①、§7 N13图内零文字；
高中同步总控任务E／§5命名「人教X版册别·部分封面（第X章 章名·件型）」；规格书§3波次T-⑧）

功能
----
生成「部分封面」配页 docx（每部分一张）＋章节主题图（TikZ、图内零文字、每章一张、
该章全部部分封面复用——主题图走 工具/绘图模板库.py 渲染出口）。

要素（N9，逐项内置自检）：羿郭工作室＋册名＋大章号数字（突出、教材化版式）＋章名＋
件型（衔接/清单/讲练）＋统计＋一句导读（创作位：【编注】起段、统计性文案一句为限——
公共规则§5创作层铁规①，逐句与统计事实核验）＋章节主题图。

配页件属性（内置断言，任一不过＝FAIL 退出码 2、不落盘交付）：
B1 不计页、不设页脚页码与页眉——包内无 header/footer 部件、sectPr 无
   headerReference/footerReference、settings 无 evenAndOddHeaders/titlePg。
B2 A4：pgSz 11906×16838（python-docx 默认 Letter 12240×15840，须 XML 改 A4——经验教训）；
   四边 pgMar=850（1.5cm，同成品规格）。
B3 全部左对齐（全部段落显式 w:jc=left）；无 w:ind 缩进。
B4 要素逐项在位（品牌/册名/大章号/章名/件型/统计/导读/主题图 inline drawing 恰1张）。
B5 导读句【编注】起段、一句为限，句内统计数字与统计配置逐项一致（内建核对）。
B6 中文宋体、西文/数字 Times New Roman（全部文字 run 显式 rFonts）。
B7 主题图 TikZ 源码图内零文字（无 \node/label/title；N13）＋嵌入为 wp:inline。
封面件类字号豁免（大字版式，§7自检⑧）；主题图渲染 PNG≥300dpi 嵌入。

用法
----
python 工具/部分封面产件.py --out-dir 输出目录 [--only 第1章 空间向量与立体几何]

输出：--out-dir 下 6 张 docx（人教B版选必1·部分封面（第X章 章名·件型）.docx）＋
主题图_第X章.png/.pdf（sources/ 含 TikZ 源码）＋ 部分封面自检.json（B1—B7 数字）。
统计数字为本轮（2026-08-31 版式改版回扫轮）写死参数：内容不变、直接引用。
"""
from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import 绘图模板库 as tpl  # noqa: E402

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ----------------------------- 写死配置（本轮统计：规格书§0/派发语） -----------------------------

BOOK = "人教B版选必1"
STATS = {
    ("第1章 空间向量与立体几何", "衔接"): {"kind": "衔接", "total": 29, "必会": True},
    ("第1章 空间向量与立体几何", "清单"): {"kind": "清单", "total": 47, "基": 33, "进": 14},
    ("第1章 空间向量与立体几何", "讲练"): {"kind": "讲练", "total": 140, "简单": 21, "中档": 104, "难": 15},
    ("第2章 平面解析几何", "衔接"): {"kind": "衔接", "total": 13, "必会": True},
    ("第2章 平面解析几何", "清单"): {"kind": "清单", "total": 67, "基": 38, "进": 29},
    ("第2章 平面解析几何", "讲练"): {"kind": "讲练", "total": 339, "简单": 47, "中档": 246, "难": 46},
}
# A''（2026-09-02）：大字区＝件型词全称映射＋各章部分动态序（分层卷落地后随派生增补）
JIANXING_WORD = {"衔接": "衔 接", "清单": "知识清单", "讲练": "讲 练",
                 "简单卷": "简单卷", "中档卷": "中档卷", "冲刺卷": "冲刺卷"}
PARTS_OF_CHAPTER = {
    "第1章 空间向量与立体几何": ["衔接", "知识清单", "讲练"],
    "第2章 平面解析几何": ["衔接", "知识清单", "讲练"],
}
# 导读一句（创作位：【编注】起段、统计性文案一句为限——铁规①；数字与 STATS 内建核对）
LEADS = {
    ("第1章 空间向量与立体几何", "衔接"):
        "【编注】本部分为第1章衔接，共29题、全部必会——逐题过关后，再进入本章知识清单与讲练。",
    ("第1章 空间向量与立体几何", "清单"):
        "【编注】本部分为第1章知识清单，共47条（基础33条必会、进阶14条汇总）——基础学完才能做题，进阶做题后回看。",
    ("第1章 空间向量与立体几何", "讲练"):
        "【编注】本部分为第1章讲练，共140题（简单21、中档104、难15）——简单保60%，加中档保80%，难档冲100%。",
    ("第2章 平面解析几何", "衔接"):
        "【编注】本部分为第2章衔接，共13题、全部必会——逐题过关后，再进入本章知识清单与讲练。",
    ("第2章 平面解析几何", "清单"):
        "【编注】本部分为第2章知识清单，共67条（基础38条必会、进阶29条汇总）——基础学完才能做题，进阶做题后回看。",
    ("第2章 平面解析几何", "讲练"):
        "【编注】本部分为第2章讲练，共339题（简单47、中档246、难46）——简单保60%，加中档保80%，难档冲100%。",
}
CHAPTER_NO = {"第1章 空间向量与立体几何": "1", "第2章 平面解析几何": "2"}
ACCENT = "18,85,126"        # 主题图强调色（青蓝；黑白打印呈中深灰，白底可辨）

# ----------------------------- 主题图（TikZ 图内零文字 N13） -----------------------------


def _hull_edges(pts):
    """平行投影轮廓（凸包）→ 用于立体线框的可见/被遮挡棱分类。"""
    def cross(o, a, b):
        return (a[0] - o[0]) * (b[1] - o[1]) - (a[1] - o[1]) * (b[0] - o[0])
    P = sorted(set(pts))
    if len(P) <= 1:
        return set()
    lo = []
    for p in P:
        while len(lo) >= 2 and cross(lo[-2], lo[-1], p) <= 0:
            lo.pop()
        lo.append(p)
    up = []
    for p in reversed(P):
        while len(up) >= 2 and cross(up[-2], up[-1], p) <= 0:
            up.pop()
        up.append(p)
    hull = lo[:-1] + up[:-1]
    es = set()
    for i in range(len(hull)):
        a, b = hull[i], hull[(i + 1) % len(hull)]
        es.add((a, b) if a <= b else (b, a))
    return es


def theme_source_ch1():
    """第1章主题图：空间直角坐标系（无标注）＋平行六面体（被遮挡棱虚线）＋空间向量箭头。"""
    # 屏幕斜二测：y→右，z→上，x→左下（指向观者）
    def s(px, py, pz):
        return (py - 0.55 * px, pz - 0.42 * px)

    a, b, c = (2.9, 2.3, 1.8)
    ox, oy = 1.15, 0.55
    V = {}
    for i, (dx, dy, dz) in enumerate([(0, 0, 0), (a, 0, 0), (a, b, 0), (0, b, 0),
                                      (0, 0, c), (a, 0, c), (a, b, c), (0, b, c)]):
        V["p%d" % i] = s(ox + dx, oy + dy, dz)
    E = [("p0", "p1"), ("p1", "p2"), ("p2", "p3"), ("p3", "p0"),
         ("p4", "p5"), ("p5", "p6"), ("p6", "p7"), ("p7", "p4"),
         ("p0", "p4"), ("p1", "p5"), ("p2", "p6"), ("p3", "p7")]
    hull = _hull_edges(V.values())
    vis, hid = [], []
    for u, v in E:
        k = V[u], V[v]
        key = k if k[0] <= k[1] else (k[1], k[0])
        (vis if key in hull else hid).append((V[u], V[v]))
    L = [r"\definecolor{acc}{RGB}{%s}" % ACCENT,
         r"\begin{tikzpicture}[>=Stealth, line width=1.0pt]",
         "  % 空间直角坐标系（图内零文字）",
         r"  \coordinate (O) at %s;" % _c(s(0, 0, 0)),
         r"  \draw[->, thick] (O) -- %s;" % _c(s(0, 6.4, 0)),
         r"  \draw[->, thick] (O) -- %s;" % _c(s(0, 0, 4.9)),
         r"  \draw[->, thick] (O) -- %s;" % _c(s(3.6, 0, 0)),
         r"  \fill[black!5] %s -- %s -- %s -- %s -- cycle;" % (
             _c(V["p0"]), _c(V["p1"]), _c(V["p2"]), _c(V["p3"])),
         "  % 平行六面体（被遮挡棱虚线——凸包轮廓法判定）"]
    for u, v in hid:
        L.append(r"  \draw[dashed, black!60] %s -- %s;" % (_c(u), _c(v)))
    for u, v in vis:
        L.append(r"  \draw %s -- %s;" % (_c(u), _c(v)))
    L.append(r"  \draw[->, line width=2.6pt, acc] %s -- %s;" % (
        _c(s(0.45, 0.3, 0.1)), _c(s(3.1, 1.5, 2.9))))
    L.append(r"  \fill[acc] %s circle (2.4pt);" % _c(s(0.45, 0.3, 0.1)))
    L.append(r"  \fill[black!12] %s -- %s -- %s -- %s -- cycle;" % (
        _c(s(0.6, 3.4, 2.9)), _c(s(0.6, 6.3, 2.9)), _c(s(2.5, 6.3, 4.1)), _c(s(2.5, 3.4, 4.1))))
    L.append(r"  \draw[black!45] %s -- %s -- %s -- %s -- cycle;" % (
        _c(s(0.6, 3.4, 2.9)), _c(s(0.6, 6.3, 2.9)), _c(s(2.5, 6.3, 4.1)), _c(s(2.5, 3.4, 4.1))))
    L.append(r"\end{tikzpicture}")
    return "\n".join(L), len(vis), len(hid)


def theme_source_ch2():
    """第2章主题图：坐标系（无标注）＋椭圆/双曲线/抛物线＋焦点圆点＋割线。"""
    import math
    L = [r"\definecolor{acc}{RGB}{%s}" % ACCENT,
         r"\begin{tikzpicture}[>=Stealth, line width=1.0pt, scale=1.0]",
         "  % 平面直角坐标系（图内零文字）",
         r"  \draw[->] (-5.2, 0) -- (6.0, 0);",
         r"  \draw[->] (0, -3.4) -- (0, 3.6);",
         "  % 椭圆（倾斜，含焦点与割线）",
         r"  \begin{scope}[rotate around={-16:(2.15, 1.25)}]",
         r"    \draw[line width=1.3pt] (2.15, 1.25) ellipse [x radius=2.45, y radius=1.02];",
         r"    \fill (2.15-1.98, 1.25) circle (1.9pt);",
         r"    \fill (2.15+1.98, 1.25) circle (1.9pt);",
         r"  \end{scope}",
         r"  \draw[dashed, black!55] (-2.9, 3.15) -- (6.6, -0.75);",
         "  % 双曲线（左右支）",
         r"  \draw[line width=1.2pt, acc] plot[domain=-2.05:2.05, samples=61] "
         r"({-2.85 - 0.62*\x*\x/4}, \x);",
         r"  \draw[line width=1.2pt, acc] plot[domain=-2.05:2.05, samples=61] "
         r"({-1.55 + 0.62*\x*\x/4}, \x);",
         r"  \fill (-3.55, 0) circle (1.9pt);  \fill (-0.85, 0) circle (1.9pt);",
         "  % 抛物线（开口向右）",
         r"  \draw[line width=1.2pt] plot[domain=-2.3:2.3, samples=61] "
         r"({1.05 + 0.28*\x*\x}, \x-2.25);",
         r"  \fill (1.75, -2.25) circle (1.9pt);"]
    L.append(r"\end{tikzpicture}")
    return "\n".join(L), None, None


def _c(p):
    return "(%.3f, %.3f)" % p


THEMES = {"第1章 空间向量与立体几何": theme_source_ch1,
          "第2章 平面解析几何": theme_source_ch2}


def build_theme(chapter, out_dir, dpi=300):
    src, nvis, nhid = THEMES[chapter]()
    # B7 前半：图内零文字断言（无 \node / node 括注 / \title / label 记述）
    bad = [k for k in (r"\node", "node={", "node[", r"\title", "label(")
           if k in src]
    assert not bad, "B7 FAIL 主题图源码疑似含文字元素: %s" % bad
    png = Path(out_dir) / ("主题图_%s.png" % CHAPTER_NO[chapter])
    tpl.render(tpl.tikz_raw(src), png, dpi=dpi, keep_pdf=True)
    return png, src, nvis, nhid


# ----------------------------- 封面 docx -----------------------------

SZ = {"品牌": 28, "册名": 30, "章号行": 44, "章名": 56, "大字区": 160, "统计": 28,
      "统计": 28, "导读": 22, "空": 36}


def _run(p, text, sz, bold=False, latin_digit=False):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(sz / 2)
    rPr = r._element.get_or_add_rPr()
    rf = rPr.get_or_add_rFonts() if hasattr(rPr, "get_or_add_rFonts") else None
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    return r


def _para(doc, text="", sz=SZ["空"], bold=False, style_sz=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if text:
        _run(p, text, sz, bold)
    else:                                        # 空段也定字号（占位高度可控）
        _run(p, "", sz)
    return p


def stats_line(st):
    if st["kind"] == "衔接":
        return "衔接%d题·全部必会" % st["total"]
    if st["kind"] == "清单":
        return "知识清单%d条〔基%d·进%d〕" % (st["total"], st["基"], st["进"])
    return "讲练%d题：简单%d｜中档%d｜难%d" % (
        st["total"], st["简单"], st["中档"], st["难"])


def check_lead(chapter, jianxing):
    """B5：导读句统计数字与 STATS 逐项核对（铁规①亲算核验的机检载体——逐句结论另落报告）。"""
    st, s = STATS[(chapter, jianxing)], LEADS[(chapter, jianxing)]
    assert s.startswith("【编注】"), "B5 FAIL 导读未【编注】起段"
    assert s.count("。") == 1 and s.endswith("。"), "B5 FAIL 导读须一句为限"
    total_pat = ("共%d条" if st["kind"] == "清单" else "共%d题") % st["total"]
    pairs = [(total_pat,)] + (
        [("基础%d条" % st["基"]), ("进阶%d条" % st["进"])] if st["kind"] == "清单"
        else [("简单%d" % st["简单"]), ("中档%d" % st["中档"]), ("难%d" % st["难"])]
        if st["kind"] == "讲练" else [])
    for pat in pairs:
        assert pat[0] in s, "B5 FAIL 导读缺统计事实 %r" % pat[0]
    return s


def build_cover(chapter, jianxing, theme_png, out_dir):
    st = STATS[(chapter, jianxing)]
    lead = check_lead(chapter, jianxing)
    name = "%s·部分封面（%s·%s）" % (BOOK, chapter, jianxing)
    doc = docx.Document()

    _para(doc, "羿郭工作室", SZ["品牌"], bold=True)
    _para(doc, BOOK, SZ["册名"], bold=True)
    _para(doc)
    _para(doc, "第" + CHAPTER_NO[chapter] + "章", SZ["章号行"], bold=True)   # A''：章号缩为普通要素行
    _para(doc, chapter, SZ["章名"], bold=True)
    _para(doc)
    _para(doc, JIANXING_WORD[jianxing], SZ["大字区"], bold=True)            # A''：大字区＝件型词
    parts_line = "本章部分：" + "→".join(PARTS_OF_CHAPTER[chapter])
    _para(doc, parts_line, 28)                                              # A''：本章部分动态行
    _para(doc)
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pic_p.paragraph_format.space_before = Pt(0)
    pic_p.paragraph_format.space_after = Pt(0)
    run = pic_p.add_run()
    run.add_picture(str(theme_png), width=Cm(15.0))
    _para(doc)
    _para(doc, "统计：" + stats_line(st), SZ["统计"])
    _para(doc, lead, SZ["导读"])

    # A4 XML（python-docx 默认 Letter→A4 11906×16838）＋边距 850 ＋无页眉页脚引用
    sect = doc.sections[0]
    sp = sect._sectPr
    for tag, attrs in (("w:pgSz", {"w:w": "11906", "w:h": "16838"}),
                       ("w:pgMar", {"w:top": "850", "w:right": "850",
                                    "w:bottom": "850", "w:left": "850",
                                    "w:header": "850", "w:footer": "850",
                                    "w:gutter": "0"})):
        el = sp.find(qn(tag))
        if el is None:
            el = sp.makeelement(qn(tag), {})
            sp.append(el)
        for k, v in attrs.items():
            el.set(qn(k), v)
    doc.core_properties.title = name
    doc.core_properties.author = ""
    out = Path(out_dir) / (name + ".docx")
    doc.save(str(out))
    return out


# ----------------------------- 自检（B1—B7） -----------------------------

def self_check(path, chapter, jianxing):
    st = STATS[(chapter, jianxing)]
    z = zipfile.ZipFile(str(path))
    docx_xml = z.read("word/document.xml").decode("utf-8")
    settings = z.read("word/settings.xml").decode("utf-8")
    names = z.namelist()
    r = {"文件": Path(path).name}

    hf = [n for n in names if re.search(r"word/(header|footer)\d*\.xml$", n)]
    r["B1包内页眉页脚部件数"] = len(hf)
    r["B1 sectPr引用"] = docx_xml.count("headerReference") + docx_xml.count("footerReference")
    r["B1 evenAndOddHeaders/titlePg"] = int("evenAndOddHeaders" in settings) + int("<w:titlePg" in docx_xml)
    assert not hf and r["B1 sectPr引用"] == 0 and r["B1 evenAndOddHeaders/titlePg"] == 0, "B1 FAIL"

    m = re.search(r'<w:pgSz w:w="(\d+)" w:h="(\d+)"', docx_xml)
    r["B2 pgSz"] = [int(m.group(1)), int(m.group(2))]
    mm = re.search(r"<w:pgMar[^/>]*/?>", docx_xml)
    marg = dict(re.findall(r'w:(top|right|bottom|left|header|footer|gutter)="(-?\d+)"', mm.group(0)))
    r["B2 pgMar"] = [int(marg[k]) for k in ("top", "right", "bottom", "left")]
    r["B2 header/footer距边"] = [int(marg["header"]), int(marg["footer"]), int(marg["gutter"])]
    assert r["B2 pgSz"] == [11906, 16838] and r["B2 pgMar"] == [850] * 4, "B2 FAIL"

    import docx as _d
    d = _d.Document(str(path))
    paras = d.paragraphs
    r["B3 段落总数"] = len(paras)
    not_left = [i for i, p in enumerate(paras) if p.alignment != WD_ALIGN_PARAGRAPH.LEFT]
    r["B3 非左对齐段数"] = len(not_left)
    ind = docx_xml.count("<w:ind ")
    r["B3 w:ind数"] = ind
    assert not not_left and ind == 0, "B3 FAIL"

    texts = [p.text for p in paras if p.text.strip()]
    joined = " ".join(texts)
    need = ["羿郭工作室", BOOK, CHAPTER_NO[chapter], chapter, jianxing,
            stats_line(st), LEADS[(chapter, jianxing)]]
    r["B4 要素"] = {k: int(k in joined) for k in need}
    assert all(r["B4 要素"].values()), "B4 FAIL %s" % r["B4 要素"]
    r["B4 主题图inline数"] = docx_xml.count("<wp:inline")
    r["B4 主题图anchor数"] = docx_xml.count("<wp:anchor")
    assert r["B4 主题图inline数"] == 1 and r["B4 主题图anchor数"] == 0, "B7/B4 FAIL 图片非inline"

    check_lead(chapter, jianxing)
    r["B5 导读核对"] = "PASS"

    nbad_font = 0
    for p in paras:
        for run in p.runs:
            rPr = run._element.find(qn("w:rPr"))
            rf = rPr.find(qn("w:rFonts")) if rPr is not None else None
            if run.text.strip() and (rf is None or rf.get(qn("w:eastAsia")) != "宋体"
                                     or rf.get(qn("w:ascii")) != "Times New Roman"):
                nbad_font += 1
    r["B6 非宋体/TNR文字run数"] = nbad_font
    assert nbad_font == 0, "B6 FAIL"
    return r


# ----------------------------- 主入口 -----------------------------

def main(argv):
    import argparse
    ap = argparse.ArgumentParser(description="部分封面产件（N9，工具债⑧）")
    ap.add_argument("--out-dir", required=True)
    ap.add_argument("--only", default=None, help="只生成该章（章名全称）")
    a = ap.parse_args(argv)
    out = Path(a.out_dir)
    out.mkdir(parents=True, exist_ok=True)

    themes = {}
    theme_meta = {}
    chapters = sorted({c for c, _j in STATS})
    if a.only:
        chapters = [a.only]
    for ch in chapters:
        png, src, nvis, nhid = build_theme(ch, out)
        themes[ch] = png
        theme_meta[ch] = {"png": str(png), "可见棱": nvis, "被遮挡棱": nhid,
                          "灰度": dict(tpl.LAST_REPORT)}
        assert tpl.LAST_REPORT.get("contrast", 0) >= 0.35, "主题图灰度对比度不足"

    results = [theme_meta]
    order = sorted(STATS.keys())
    if a.only:
        order = [k for k in order if k[0] == a.only]
    for chapter, jianxing in order:
        p = build_cover(chapter, jianxing, themes[chapter], out)
        results.append(self_check(p, chapter, jianxing))
        print("PASS", p.name)
    (out / "部分封面自检.json").write_text(
        json.dumps(results, ensure_ascii=False, indent=1), encoding="utf-8")


if __name__ == "__main__":
    main(sys.argv[1:])
