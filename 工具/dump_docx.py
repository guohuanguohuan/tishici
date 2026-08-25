# -*- coding: utf-8 -*-
"""dump_docx.py — docx 全文转储（按文档顺序，OMML 线性化保留结构分隔符）
用法:
  python dump_docx.py <docx路径> <输出txt路径>          # 全量转储（原行为）
  python dump_docx.py --index <docx路径> <输出md路径>    # 题块索引：题号/元素区间/首句/标签/图数（会话先读索引再定点切片，省上下文）
  python dump_docx.py --slice <docx路径> 起:末 <输出txt> # 只转储体元素序号[起,末]（含端点）——按索引的区间定点读取
  python dump_docx.py --indexdir <目录> <输出md路径>     # 目录级批量普查草稿（文件/题块数/图数/首句/字节，供素材普查档案回填）
元素序号＝body 直接子元素（段落与表格都计数）的 0 基序号，索引与切片共用同一序号体系。
可复用工具：数学物理单元同步（题目台账/亲算/修复轮定位用）"""
import sys, re, os
from docx import Document
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WC, MC = '{'+W+'}', '{'+M+'}'

# 题块起始判定：全角「N．」开头且该段长度≥8（经验口径：小数/年份不起误报）
QNUM_RE = re.compile(r'^(\d{1,3})．')
QNUM_MINLEN = 8
# 讲练一体栏目起始（也算块边界）
MARK_RE = re.compile(r'^(【例题】|【典例\d*】|【举一反三】|【练习题】|【易错题)')
# 结构/节标题行（索引里单独列出便于导航）
HEAD_RE = re.compile(r'^\d+(\.\d+)+\s*\S')

def tagof(el):
    return el.tag.split('}')[-1]

def omml_text(el):
    """递归线性化 OMML，保留 m:d 分隔符 / m:rad 根号 / m:f 分数 / 上下标"""
    t = tagof(el)
    if t == 't':
        return el.text or ''
    if t == 'f':  # 分数
        num = el.find(MC+'num'); den = el.find(MC+'den')
        return '(' + ''.join(omml_text(c) for c in num) + ')/(' + ''.join(omml_text(c) for c in den) + ')'
    if t == 'rad':
        deg = None; e = el.find(MC+'e')
        d = el.find(MC+'deg')
        if d is not None:
            dt = ''.join(omml_text(c) for c in d)
            deg = dt.strip()
        inner = ''.join(omml_text(c) for c in e) if e is not None else ''
        return ('root[' + deg + '](' if deg else '√(') + inner + ')'
    if t == 'sSup':
        e = el.find(MC+'e'); sup = el.find(MC+'sup')
        return ''.join(omml_text(c) for c in e) + '^(' + ''.join(omml_text(c) for c in sup) + ')'
    if t == 'sSub':
        e = el.find(MC+'e'); sub = el.find(MC+'sub')
        return ''.join(omml_text(c) for c in e) + '_(' + ''.join(omml_text(c) for c in sub) + ')'
    if t == 'sPre':
        sub = el.find(MC+'sub'); sup = el.find(MC+'sup'); e = el.find(MC+'e')
        return '_(' + (''.join(omml_text(c) for c in sub) if sub is not None else '') + ')^(' + \
               (''.join(omml_text(c) for c in sup) if sup is not None else '') + ')' + \
               (''.join(omml_text(c) for c in e) if e is not None else '')
    if t == 'd':  # 括号/绝对值等定界符
        fPr = el.find(MC+'dPr')
        beg, end = '(', ')'
        if fPr is not None:
            b = fPr.find(MC+'begChr'); e2 = fPr.find(MC+'endChr')
            if b is not None: beg = b.get(MC+'val') or ''
            if e2 is not None: end = e2.get(MC+'val') or ''
            if b is None: beg = '('
            if e2 is None: end = ')'
        inner = ''
        for e3 in el.findall(MC+'e'):
            inner += ''.join(omml_text(c) for c in e3) + ','
        return beg + inner.rstrip(',') + end
    if t == 'nary':  # 求和/乘积
        sub = el.find(MC+'sub'); sup = el.find(MC+'sup'); e = el.find(MC+'e')
        chrEl = el.find(MC+'naryPr')
        op = '∑'
        if chrEl is not None:
            c = chrEl.find(MC+'chr')
            if c is not None and c.get(MC+'val'): op = c.get(MC+'val')
        return op + '_(' + (''.join(omml_text(c) for c in sub) if sub is not None else '') + ')^(' + \
               (''.join(omml_text(c) for c in sup) if sup is not None else '') + ')' + \
               (''.join(omml_text(c) for c in e) if e is not None else '')
    if t == 'func':
        fname = el.find(MC+'fName'); e = el.find(MC+'e')
        return (''.join(omml_text(c) for c in fname) if fname is not None else '') + \
               ('(' + ''.join(omml_text(c) for c in e) + ')' if e is not None else '')
    if t in ('r',):
        txt = ''
        for c in el:
            ct = tagof(c)
            if ct == 't': txt += c.text or ''
        return txt
    if t == 'eqArr':
        parts = []
        for e4 in el.findall(MC+'e'):
            parts.append(''.join(omml_text(c) for c in e4))
        return ' ‖ '.join(parts)
    # 其余容器：递归
    out = ''
    for c in el:
        ct = tagof(c)
        if ct.endswith('Pr'): continue
        out += omml_text(c)
    return out

def para_text(p_el):
    """段落全文（w:t + m:t 按文档顺序）"""
    parts = []
    def walk(el):
        for c in el:
            ct = tagof(c)
            if ct == 't' and c.tag.startswith(WC):
                parts.append(c.text or '')
            elif ct == 'oMath':
                parts.append('⟦' + omml_text(c) + '⟧')
            elif ct == 'oMathPara':
                for om in c.findall(MC+'oMath'):
                    parts.append('⟦' + omml_text(om) + '⟧')
            elif ct in ('drawing', 'pict', 'object'):
                parts.append('【图】')
            elif ct == 'br':
                parts.append('⏎')
            elif ct == 'tab':
                parts.append('\t')
            else:
                walk(c)
    walk(p_el)
    return ''.join(parts)

def body_elements(src):
    """返回 [(元素序号, tag, 段落文本或None)]，表格计数为独立元素。"""
    doc = Document(src)
    out = []
    for i, child in enumerate(doc.element.body.iterchildren()):
        ct = tagof(child)
        if ct == 'p':
            out.append((i, 'p', para_text(child)))
        elif ct == 'tbl':
            out.append((i, 'tbl', None))
        elif ct == 'sectPr':
            out.append((i, 'sectPr', None))
        else:
            out.append((i, ct, None))
    return out

def is_qstart(text):
    if not text:
        return None
    t = text.strip()
    if len(t) >= QNUM_MINLEN:
        m = QNUM_RE.match(t)
        if m:
            return m.group(1)
    if MARK_RE.match(t):
        return t[:8]
    return None

def md_safe(s):
    """markdown 表格安全化：竖线换全角（OMML 绝对值线性化会产生 |）。"""
    return (s or '').replace('|', '｜')

def label_val(block_text, name, maxlen=24):
    m = re.search('【'+name+'】\\s*([^\\n【]{0,%d})' % maxlen, block_text)
    return md_safe(m.group(1).strip() if m else '')

def make_index(src, dst):
    els = body_elements(src)
    starts = []   # (元素序号, 题号/栏目名, 首句)
    for i, tag, text in els:
        if tag != 'p':
            continue
        q = is_qstart(text)
        if q is not None:
            first = text.strip().replace('|', '／')
            starts.append((i, q, first[:40]))
    # 块区间＝起点到下一起点前一个元素（无下一起点则到文末）
    rows = []
    for k, (i, q, first) in enumerate(starts):
        j = starts[k+1][0] - 1 if k + 1 < len(starts) else els[-1][0]
        block = '\n'.join(t for ii, tag, t in els if i <= ii <= j and t is not None)
        rows.append('| {q} | {a}-{b} | {first} | {ans} | {diff} | {kp} | {img} | {det} |'.format(
            q=q, a=i, b=j, first=first,
            ans=label_val(block, '答案'), diff=label_val(block, '难度'),
            kp=label_val(block, '知识点'), img=block.count('【图】'),
            det='有' if '【详解】' in block else ''))
    # 结构/节标题导航行
    heads = ['| {a} | {t} |'.format(a=i, t=text.strip().replace('|', '／')[:36])
             for i, tag, text in els if tag == 'p' and text and HEAD_RE.match(text.strip()) and len(text.strip()) < 40]
    with open(dst, 'w', encoding='utf-8') as f:
        f.write('# 题块索引：%s\n\n' % src)
        f.write('> 用法：先读本表定位（元素区间），再 `dump_docx.py --slice 原docx 起:末 输出.txt` 定点读取；'
                '标签/图数由文本层提取，仅供导航，判定以原文为准。\n\n')
        f.write('| 题号/栏目 | 元素区间 | 首句(40字) | 答案 | 原难度 | 原知识点 | 图数 | 详解 |\n|---|---|---|---|---|---|---|---|\n')
        f.write('\n'.join(rows))
        f.write('\n\n## 结构/节标题行\n\n| 元素序号 | 文本 |\n|---|---|\n')
        f.write('\n'.join(heads))
    print('OK index %s -> %s  题块=%d 标题行=%d' % (src, dst, len(rows), len(heads)))

def make_slice(src, rng, dst):
    a, b = rng.split(':')
    a, b = int(a), int(b)
    doc = Document(src)
    lines = []
    tbl_i = 0
    for i, child in enumerate(doc.element.body.iterchildren()):
        if i < a or i > b:
            continue
        ct = tagof(child)
        if ct == 'p':
            lines.append('[%d] %s' % (i, para_text(child)))
        elif ct == 'tbl':
            tbl_i += 1
            lines.append('[%d] ───表格%d───' % (i, tbl_i))
            for tr in child.findall(WC+'tr'):
                cells = []
                for tc in tr.findall(WC+'tc'):
                    ctxt = ' '.join(para_text(p) for p in tc.findall('.//'+WC+'p'))
                    cells.append(ctxt.strip())
                lines.append(' | '.join(cells))
    with open(dst, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print('OK slice %s [%d:%d] -> %s  行数=%d' % (src, a, b, dst, len(lines)))

def make_indexdir(root, dst):
    """目录级批量普查草稿：文件｜题块数｜图数｜首块首句｜字节数（供素材普查档案回填）。"""
    rows = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames.sort()
        for fn in sorted(filenames):
            if not fn.lower().endswith('.docx') or fn.startswith('~$'):
                continue
            p = os.path.join(dirpath, fn)
            rel = os.path.relpath(p, root).replace('\\', '／')
            try:
                els = body_elements(p)
                qs, imgs, first = 0, 0, ''
                for i, tag, text in els:
                    if tag != 'p' or text is None:
                        continue
                    if is_qstart(text) is not None:
                        qs += 1
                        if not first:
                            first = text.strip()[:30].replace('|', '／')
                    imgs += text.count('【图】')
                rows.append('| %s | %d | %d | %s | %d |' % (rel, qs, imgs, first, os.path.getsize(p)))
            except Exception as e:
                rows.append('| %s | 错误 | | %s | %d |' % (rel, str(e)[:24], os.path.getsize(p)))
    with open(dst, 'w', encoding='utf-8') as f:
        f.write('# 批量普查草稿：%s\n\n' % root)
        f.write('> 由 工具/dump_docx.py --indexdir 生成；供 素材普查/素材普查-*.md 回填——'
                '机械列直接取用，结构坑位与归属两列须开卷/按文件夹人工补。\n\n')
        f.write('| 相对路径 | 题块数 | 图数 | 首块首句(30字) | 字节数 |\n|---|---|---|---|---|\n')
        f.write('\n'.join(rows))
    print('OK indexdir %s -> %s  文件数=%d' % (root, dst, len(rows)))

def main():
    args = sys.argv[1:]
    if args and args[0] == '--index':
        make_index(args[1], args[2])
    elif args and args[0] == '--slice':
        make_slice(args[1], args[2], args[3])
    elif args and args[0] == '--indexdir':
        make_indexdir(args[1], args[2])
    else:
        src, dst = args[0], args[1]
        doc = Document(src)
        body = doc.element.body
        lines = []
        tbl_i = 0
        for child in body.iterchildren():
            ct = tagof(child)
            if ct == 'p':
                t = para_text(child)
                lines.append(t)
            elif ct == 'tbl':
                tbl_i += 1
                lines.append(f'───表格{tbl_i}───')
                for tr in child.findall(WC+'tr'):
                    cells = []
                    for tc in tr.findall(WC+'tc'):
                        ctxt = ' '.join(para_text(p) for p in tc.findall('.//'+WC+'p'))
                        cells.append(ctxt.strip())
                    lines.append(' | '.join(cells))
                lines.append(f'───表格{tbl_i}结束───')
        with open(dst, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        n = len([l for l in lines if l.strip()])
        print(f'OK {src} -> {dst}  非空行数={n} 表格数={tbl_i}')

if __name__ == '__main__':
    main()
