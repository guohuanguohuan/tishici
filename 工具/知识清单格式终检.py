# -*- coding: utf-8 -*-
"""知识清单 DOCX 的确定性格式终检与修复。

⚠ 标记方案警示（2026-08-25 起）：标记方案已改为「仅方框 w:bdr」（不加粗不划线），
本工具的下划线加粗逻辑（set_thick_underline / reinforce_math_emphasis /
narrow_explanatory_underlines）属旧方案遗留——只可用于处理历史旧件的格式统一，
且处理旧件的正确顺序是先跑本工具再跑 标记转方框.py；对已是新方案的件（含 w:bdr）
运行本工具会加回粗下划线造成方案回退，禁止。

用途：保留原文档结构、图片和 OMML 公式，只统一 A4/页边距/段落/字体，
确保已有选择性下划线同时加粗，并把块级公式安全转为行内公式。
"""
from __future__ import annotations

import copy
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def ensure(parent, tag):
    child = parent.find(tag)
    if child is None:
        child = OxmlElement(tag.replace(W, "w:").replace(M, "m:"))
        parent.append(child)
    return child


def set_bool(rpr, local):
    el = rpr.find(W + local)
    if el is None:
        el = OxmlElement("w:" + local)
        rpr.append(el)
    el.set(W + "val", "1")


def set_thick_underline(rpr):
    el = rpr.find(W + "u")
    if el is None:
        el = OxmlElement("w:u")
        rpr.append(el)
    el.set(W + "val", "thick")


def has_pua(text):
    return any("\ue000" <= ch <= "\uf8ff" for ch in text)


def format_text_run(run_el):
    text = "".join(run_el.itertext())
    rpr = run_el.find(W + "rPr")
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_el.insert(0, rpr)
    if not has_pua(text):
        fonts = rpr.find(W + "rFonts")
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        fonts.set(W + "ascii", "Times New Roman")
        fonts.set(W + "hAnsi", "Times New Roman")
        fonts.set(W + "eastAsia", "宋体")
        fonts.set(W + "cs", "Times New Roman")
    for name in ("sz", "szCs"):
        el = rpr.find(W + name)
        if el is None:
            el = OxmlElement("w:" + name)
            rpr.append(el)
        el.set(W + "val", "21")
    if rpr.find(W + "u") is not None:
        set_bool(rpr, "b")
        set_bool(rpr, "bCs")
        set_thick_underline(rpr)


def format_paragraph(p_el):
    ppr = p_el.find(W + "pPr")
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p_el.insert(0, ppr)
    jc = ppr.find(W + "jc")
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(W + "val", "left")
    spacing = ppr.find(W + "spacing")
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(W + "before", "0")
    spacing.set(W + "after", "0")
    spacing.set(W + "line", "300")
    spacing.set(W + "lineRule", "auto")
    for name in ("pageBreakBefore", "keepNext", "keepLines"):
        el = ppr.find(W + name)
        if el is not None:
            ppr.remove(el)
    for br in p_el.findall(".//" + W + "br"):
        if br.get(W + "type") == "page":
            br.getparent().remove(br)
    marker_rpr = ppr.find(W + "rPr")
    if marker_rpr is None:
        marker_rpr = OxmlElement("w:rPr")
        ppr.append(marker_rpr)
    for name in ("sz", "szCs"):
        el = marker_rpr.find(W + name)
        if el is None:
            el = OxmlElement("w:" + name)
            marker_rpr.append(el)
        el.set(W + "val", "21")
    for run_el in p_el.findall(".//" + W + "r"):
        format_text_run(run_el)


def convert_omathpara(root):
    count = 0
    for op in list(root.findall(".//" + M + "oMathPara")):
        parent = op.getparent()
        idx = parent.index(op)
        maths = op.findall(M + "oMath")
        for math in maths:
            parent.insert(idx, copy.deepcopy(math))
            idx += 1
        parent.remove(op)
        count += 1
    return count


def reinforce_math_emphasis(root):
    for mr in root.findall(".//" + M + "r"):
        rpr = mr.find(W + "rPr")
        if rpr is not None and rpr.find(W + "u") is not None:
            set_bool(rpr, "b")
            set_bool(rpr, "bCs")
            set_thick_underline(rpr)
    for frac in root.findall(".//" + M + "f"):
        if not frac.findall(".//" + W + "u"):
            continue
        fpr = frac.find(M + "fPr")
        if fpr is None:
            fpr = OxmlElement("m:fPr")
            frac.insert(0, fpr)
        ctrl = fpr.find(M + "ctrlPr")
        if ctrl is None:
            ctrl = OxmlElement("m:ctrlPr")
            fpr.append(ctrl)
        rpr = ctrl.find(W + "rPr")
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            ctrl.append(rpr)
        set_bool(rpr, "b")
        set_bool(rpr, "bCs")
        set_thick_underline(rpr)


def replace_run_with_segments(run_el, segments):
    parent = run_el.getparent()
    idx = parent.index(run_el)
    for text, emphasized in segments:
        new_run = copy.deepcopy(run_el)
        text_nodes = new_run.findall(".//" + W + "t")
        if not text_nodes:
            continue
        text_nodes[0].text = text
        for extra in text_nodes[1:]:
            extra.getparent().remove(extra)
        rpr = new_run.find(W + "rPr")
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            new_run.insert(0, rpr)
        if emphasized:
            set_bool(rpr, "b")
            set_bool(rpr, "bCs")
            set_thick_underline(rpr)
        else:
            for name in ("b", "bCs", "u"):
                el = rpr.find(W + name)
                if el is not None:
                    rpr.remove(el)
        parent.insert(idx, new_run)
        idx += 1
    parent.remove(run_el)


def narrow_explanatory_underlines(doc):
    mappings = {
        "8．": [
            ("一个排列是指从n个不同元素中，任取m个元素，按照一定顺序排成一列的", False),
            ("一种具体排法", True), ("，它", False), ("不是数", True),
            ("；而排列数是指从n个不同元素中取m个不同元素的", False),
            ("所有排列的个数", True), ("，它是", False), ("一个数", True),
        ],
        "12．": [
            ("联系：二者都是从n个不同的元素中取出m（m≤n）个元素；区别：", False),
            ("排列与元素的顺序有关", True), ("，", False),
            ("组合与元素的顺序无关", True),
        ],
        "13．": [
            ("“", False), ("组合", True),
            ("”是指“从n个不同元素中取m（m≤n）个元素作为一组”，它不是一个数，而是", False),
            ("具体的一件事", True), ("；“", False), ("组合数", True),
            ("”是指“从n个不同元素中取出m（m≤n）个元素的", False),
            ("所有不同组合的个数", True), ("”，它是", False), ("一个数", True),
        ],
    }
    for p in doc.paragraphs:
        key = next((k for k in mappings if p.text.startswith(k)), None)
        if key is None:
            continue
        expected = "".join(text for text, _ in mappings[key])
        candidates = []
        for run in p.runs:
            rpr = run._r.find(W + "rPr")
            if rpr is not None and rpr.find(W + "u") is not None:
                candidates.append(run)
        if len(candidates) == 1 and candidates[0].text == expected:
            replace_run_with_segments(candidates[0]._r, mappings[key])


def rebuild_one_footer(footer):
    p = footer.paragraphs[0]
    for child in list(p._p):
        if child.tag != W + "pPr":
            p._p.remove(child)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_begin, instr, fld_end))
    format_text_run(run._r)


def rebuild_footer(section):
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    rebuild_one_footer(section.footer)
    rebuild_one_footer(section.even_page_footer)
    rebuild_one_footer(section.first_page_footer)


def make_title_bold(doc, expected_title):
    first = next((p for p in doc.paragraphs if p.text.strip()), None)
    if first is None:
        raise ValueError("文档没有非空标题段")
    if first.text.strip() != expected_title:
        raise ValueError(f"文内标题不匹配：{first.text!r} != {expected_title!r}")
    for run in first.runs:
        run.bold = True


def remove_trailing_empty_paragraphs(doc):
    """删除正文末尾会单独挤出空白页的纯空段。"""
    body = doc.element.body
    removed = 0
    while len(body) >= 2:
        candidate = body[-2] if body[-1].tag == W + "sectPr" else body[-1]
        if candidate.tag != W + "p":
            break
        has_content = (
            any((t.text or "").strip() for t in candidate.findall(".//" + W + "t"))
            or bool(candidate.findall(".//" + W + "drawing"))
            or bool(candidate.findall(".//" + W + "pict"))
            or bool(candidate.findall(".//" + M + "oMath"))
        )
        if has_content:
            break
        body.remove(candidate)
        removed += 1
    return removed


def process(path_str):
    path = Path(path_str).resolve()
    expected_title = path.stem
    doc = Document(str(path))
    make_title_bold(doc, expected_title)
    narrow_explanatory_underlines(doc)
    doc.core_properties.title = expected_title
    doc.core_properties.author = ""
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
        section.header_distance = Mm(15)
        section.footer_distance = Mm(15)
        section.gutter = Mm(0)
        rebuild_footer(section)
    root = doc.element
    converted = convert_omathpara(root)
    trailing_removed = remove_trailing_empty_paragraphs(doc)
    for p_el in root.findall(".//" + W + "p"):
        format_paragraph(p_el)
    reinforce_math_emphasis(root)
    tmp = path.with_name(path.stem + ".tmp.docx")
    doc.save(str(tmp))
    os.replace(tmp, path)
    print(f"OK {path} 块级公式转行内={converted} 尾部空段删除={trailing_removed}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        raise SystemExit("用法：python 知识清单格式终检.py <docx> [docx...]")
    for item in sys.argv[1:]:
        process(item)
